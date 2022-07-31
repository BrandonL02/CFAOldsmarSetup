import os

import datetime as dt

import calendar

import random
#Pull employee number from CFA_Employees.py
from CFA_Employees import Employee_num

#Returm employee number to name 
from CFA_Employees import Employee_name

#Open document for editing
from docx import Document
from docx.shared import Pt 

doc = Document('C:/Users/Brand/Downloads/Front Counter Setup Code.docx')

tables = doc.tables 
#Get date and weekday within python
date = dt.date.today()

print_date = (calendar.day_name[date.weekday()] + ', ' + date.strftime("%B %d"))
#Prompt user to enter scheduled employee quantity 
employee_count = int(input('How many employees are scheduled between 2 pm and 10:45 pm? '))
#Prompt user to enter employee names and shifts in format: FirstName LastName.StartTime.Endtime 
print("\nPlease enter who is scheduled to work tonight along with their start and end times (in military time) separated by '.'")
#Define Lists
employee_and_shift = []
employee_list = []
start_shift_list = []
end_shift_list = []
shift_length = []
time_2to3 = []
time_3to4 = []
time_4to5 = []
time_5to6 = []
time_6to7 = []
time_7to8 = []
time_8to9 = []
time_9to10 = []
FOHCrew = []

#Define positions as variables
#2-3
window2to3 = ()
DTbag2to3 = ()
Drinks2to3 = ()
Desserts2to3 = ()
Sauce2to3 = ()
FCbag2to3 = ()
Register1_2to3 = ()
Register2_2to3 = ()
Dining2to3 = ()
Cash2to3 = ()
Ipad1_2to3 = ()
Ipad2_2to3 = ()
Ipad3_2to3 = ()

#3-4
window3to4 = ()
DTbag3to4 = ()
Drinks3to4 = ()
Desserts3to4 = ()
Sauce3to4 = ()
FCbag3to4 = ()
Register1_3to4 = ()
Register2_3to4 = ()
Dining3to4 = ()
Cash3to4 = ()
Ipad1_3to4 = ()
Ipad2_3to4 = ()
Ipad3_3to4 = ()

#4-5
window4to5 = ()
DTbag4to5 = ()
Drinks4to5 = ()
Desserts4to5 = ()
Sauce4to5 = ()
FCbag4to5 = ()
Register1_4to5 = ()
Register2_4to5 = ()
Dining4to5 = ()
Cash4to5 = ()
Ipad1_4to5 = ()
Ipad2_4to5 = ()
Ipad3_4to5 = ()

#5-6
window5to6 = ()
DTbag5to6 = ()
Drinks5to6 = ()
Desserts5to6 = ()
Sauce5to6 = ()
FCbag5to6 = ()
Register1_5to6 = ()
Register2_5to6 = ()
Dining5to6 = ()
Cash5to6 = ()
Ipad1_5to6 = ()
Ipad2_5to6 = ()
Ipad3_5to6 = ()

#6-7
window6to7 = ()
DTbag6to7 = ()
Drinks6to7 = ()
Desserts6to7 = ()
Sauce6to7 = ()
FCbag6to7 = ()
Register1_6to7 = ()
Register2_6to7 = ()
Dining6to7 = ()
Cash6to7 = ()
Ipad1_6to7 = ()
Ipad2_6to7 = ()
Ipad3_6to7 = ()

#7-8
window7to8 = ()
DTbag7to8 = ()
Drinks7to8 = ()
Desserts7to8 = ()
Sauce7to8 = ()
FCbag7to8 = ()
Register1_7to8 = ()
Register2_7to8 = ()
Dining7to8 = ()
Cash7to8 = ()
Ipad1_7to8 = ()
Ipad2_7to8 = ()
Ipad3_7to8 = ()

#8-9
window8to9 = ()
DTbag8to9 = ()
Drinks8to9 = ()
Desserts8to9 = ()
Sauce8to9 = ()
FCbag8to9 = ()
Register1_8to9 = ()
Register2_8to9 = ()
Dining8to9 = ()
Cash8to9 = ()
Ipad1_8to9 = ()
Ipad2_8to9 = ()
Ipad3_8to9 = ()

#9-10
window9to10 = ()
DTbag9to10 = ()
Drinks9to10 = ()
Desserts9to10 = ()
Sauce9to10 = ()
FCbag9to10 = ()
Register1_9to10 = ()
Register2_9to10 = ()
Dining9to10 = ()

pos_Ipad1_4to5 = ()
pos_Ipad2_4to5 = ()
pos_Ipad3_4to5 = ()


#Create list of scheduled employees and their shifts by repeating as many times as there are sheduled employees
full_staff = input ()
scheduled_shift = (full_staff.split('-'))
scheduled_shift = list (scheduled_shift)

for x in scheduled_shift:
    employee, start_shift, end_shift = (x.split('.',3))
    #Replace provided employee name with their given number
    FOHCrew.append(employee)
    if employee in Employee_num:
        employee = Employee_num[employee]    
    employee_list.append (employee)
    #Calculate shift length
    if (int(end_shift) - int(start_shift)) > 0:
        shift_length.append (int(end_shift) - int(start_shift))
    else:
        shift_length.append ((12 - int(start_shift)) + int(end_shift))
    if (int (end_shift) > 14 and int(start_shift) <= 14):  
        time_2to3.append(employee)
    if (int (end_shift) > 15 and int(start_shift) <= 15):  
        time_3to4.append(employee) 
    if (int (end_shift) > 16 and int(start_shift) <= 16):  
        time_4to5.append(employee)
    if (int (end_shift) > 17 and int(start_shift) <= 17):  
        time_5to6.append(employee)
    if (int (end_shift) > 18 and int(start_shift) <= 18):  
        time_6to7.append(employee)
    if (int (end_shift) > 19 and int(start_shift) <= 19):  
        time_7to8.append(employee)
    if (int (end_shift) > 20 and int(start_shift) <= 20):  
        time_8to9.append(employee)
    if (int (end_shift) > 21 and int(start_shift) <= 21):  
        time_9to10.append(employee)       
    start_shift_list.append (start_shift)
    end_shift_list.append (end_shift)
    employee_and_shift.append(scheduled_shift)
#Convert all lists of strings to integers
employee_list = list(map(int,employee_list))    
start_shift_list= list(map(int,start_shift_list))
end_shift_list = list(map(int,end_shift_list))

#Breaks

#2-3
print ("\nFor 2 pm to 3 pm:")
for x in time_2to3:
    while time_2to3 != []:
    #Register 1
        if '42' in time_3to4:
           print("Taylor on Register")
           time_2to3.remove('42') 
        else:
            pos_Register1_2to3 = random.choice(time_2to3)
            if pos_Register1_2to3 in Employee_name:
             Register1_2to3 = Employee_name[pos_Register1_2to3]
             time_2to3.remove(pos_Register1_2to3)
             print((Register1_2to3) + " on Register") 
        if time_2to3 == []:
             continue          
    #Window    
        pos_window2to3 = random.choice(time_2to3)
        if pos_window2to3 in Employee_name:
           window2to3 = Employee_name[pos_window2to3]
        time_2to3.remove(pos_window2to3)
        print((window2to3) + " on Window")
        if time_2to3 == []:
            continue
    #DT bagging    
        pos_DTbag2to3 = random.choice(time_2to3)
        if pos_DTbag2to3 in Employee_name:
           DTbag2to3 = Employee_name[pos_DTbag2to3]
        time_2to3.remove(pos_DTbag2to3)
        print((DTbag2to3) + " on DT Bagging")
        if time_2to3 == []:
            continue
    #Drinks
        pos_Drinks2to3 = random.choice(time_2to3)
        if pos_Drinks2to3 in Employee_name:
           Drinks2to3 = Employee_name[pos_Drinks2to3]
        time_2to3.remove(pos_Drinks2to3)
        print((Drinks2to3) + " on Drinks")
        if time_2to3 == []:
            continue
    #IPad 1
        pos_Ipad1_2to3 = ()
        pos_Ipad2_2to3 = ()
        pos_Ipad3_2to3 = ()
        pos_Ipad1_2to3 = random.choice(time_2to3)
        if pos_Ipad1_2to3 in Employee_name:
           Ipad1_2to3 = Employee_name[pos_Ipad1_2to3]
        time_2to3.remove(pos_Ipad1_2to3)
        print((Ipad1_2to3) + " on IPad")
        if time_2to3 == []:
            continue
    #Cash
        pos_Cash2to3 = random.choice(time_2to3)
        if pos_Cash2to3 in Employee_name:
           Cash2to3 = Employee_name[pos_Cash2to3]
        time_2to3.remove(pos_Cash2to3)
        print((Cash2to3) + " on Cash")
        if time_2to3 == []:
            continue
    #FC bagging
        pos_FCbag2to3 = random.choice(time_2to3)
        if pos_FCbag2to3 in Employee_name:
           FCbag2to3 = Employee_name[pos_FCbag2to3]
        time_2to3.remove(pos_FCbag2to3)
        print((FCbag2to3) + " on FC Bagging")
        if time_2to3 == []:
            continue
    #IPad 2
        pos_Ipad2_2to3 = random.choice(time_2to3)
        if pos_Ipad2_2to3 in Employee_name:
           Ipad2_2to3 = Employee_name[pos_Ipad2_2to3]
        time_2to3.remove(pos_Ipad2_2to3)
        print((Ipad2_2to3) + " on IPad")
        if time_2to3 == []:
            continue
    #Desserts
        pos_Desserts2to3 = random.choice(time_2to3)
        if pos_Desserts2to3 in Employee_name:
           Desserts2to3 = Employee_name[pos_Desserts2to3]
        time_2to3.remove(pos_Desserts2to3)
        print((Desserts2to3) + " on Desserts")
        if time_2to3 == []:
            continue
    #Register 2
        pos_Register2_2to3 = random.choice(time_2to3)
        if pos_Register2_2to3 in Employee_name:
           Register2_2to3 = Employee_name[pos_Register2_2to3]
        time_2to3.remove(pos_Register2_2to3)
        print((Register2_2to3) + " on Register")
        if time_2to3 == []:
            continue
    #IPad 3
        pos_Ipad3_2to3 = random.choice(time_2to3)
        if pos_Ipad3_2to3 in Employee_name:
           Ipad3_2to3 = Employee_name[pos_Ipad3_2to3]
        time_2to3.remove(pos_Ipad3_2to3)
        print((Ipad3_2to3) + " on IPad")
        if time_2to3 == []:
            continue
    #Dining Room
        pos_Dining2to3 = random.choice(time_2to3)
        if pos_Dining2to3 in Employee_name:
           Dining2to3 = Employee_name[pos_Dining2to3]
        time_2to3.remove(pos_Dining2to3)
        print((Dining2to3) + " on Dining Room")
        if time_2to3 == []:
            continue
    #Sauces
        pos_Sauce2to3 = random.choice(time_2to3)
        if pos_Sauce2to3 in Employee_name:
           Sauce2to3 = Employee_name[pos_Sauce2to3]
        time_2to3.remove(pos_Sauce2to3)
        print((Sauce2to3) + " on Sauces") 
        if time_2to3 == []:
            break
#3 to 4
print ("\nFor 3 pm to 4 pm:")
for x in time_3to4:
    while time_3to4 != []:
    #Register 1
        if '42' in time_3to4:
           print("Taylor on Register")
           time_3to4.remove('42') 
        else:
            pos_Register1_3to4 = random.choice(time_3to4)
            if pos_Register1_3to4 in Employee_name:
             Register1_3to4 = Employee_name[pos_Register1_3to4]
             time_3to4.remove(pos_Register1_3to4)
             print((Register1_3to4) + " on Register") 
        if time_3to4 == []:
             continue     
#Window    
        pos_window3to4 = random.choice(time_3to4)
        if pos_window3to4 in Employee_name:
           window3to4 = Employee_name[pos_window3to4]
        time_3to4.remove(pos_window3to4)
        print((window3to4) + " on Window")
        if time_3to4 == []:
            continue
    #DT bagging    
        pos_DTbag3to4 = random.choice(time_3to4)
        if pos_DTbag3to4 in Employee_name:
           DTbag3to4 = Employee_name[pos_DTbag3to4]
        time_3to4.remove(pos_DTbag3to4)
        print((DTbag3to4) + " on DT Bagging")
        if time_3to4 == []:
            continue
    #Drinks
        pos_Drinks3to4 = random.choice(time_3to4)
        if pos_Drinks3to4 in Employee_name:
           Drinks3to4 = Employee_name[pos_Drinks3to4]
        time_3to4.remove(pos_Drinks3to4)
        print((Drinks3to4) + " on Drinks")
        if time_3to4 == []:
            continue
    #IPad 1
        pos_Ipad1_3to4 = ()
        pos_Ipad2_3to4 = ()
        pos_Ipad3_3to4 = ()
        pos_Ipad1_3to4 = random.choice(time_3to4)
        while pos_Ipad1_3to4 == pos_Ipad1_2to3 or pos_Ipad1_3to4 == pos_Ipad2_2to3 or pos_Ipad1_3to4 == pos_Ipad3_2to3:
            if len(time_3to4) == 1:
                 pos_Ipad1_3to4 = ()   
                 break
            else:
                pos_Ipad1_3to4 = random.choice(time_3to4)
        if pos_Ipad1_3to4 in Employee_name:
           Ipad1_3to4 = Employee_name[pos_Ipad1_3to4]
           time_3to4.remove(pos_Ipad1_3to4)
           print((Ipad1_3to4) + " on IPad")
        if time_3to4 == []:
            continue
    #Cash
        pos_Cash3to4 = random.choice(time_3to4)
        if pos_Cash3to4 in Employee_name:
           Cash3to4 = Employee_name[pos_Cash3to4]
        time_3to4.remove(pos_Cash3to4)
        print((Cash3to4) + " on Cash")
        if time_3to4 == []:
            continue
    #FC bagging
        pos_FCbag3to4 = random.choice(time_3to4)
        if pos_FCbag3to4 in Employee_name:
           FCbag3to4 = Employee_name[pos_FCbag3to4]
        time_3to4.remove(pos_FCbag3to4)
        print((FCbag3to4) + " on FC Bagging")
        if time_3to4 == []:
            continue
    #IPad 2
        pos_Ipad2_3to4 = random.choice(time_3to4)
        while pos_Ipad2_3to4 == pos_Ipad1_2to3 or pos_Ipad2_3to4 == pos_Ipad2_2to3 or pos_Ipad2_3to4 == pos_Ipad3_2to3:   
            if len(time_3to4) == 1:
                pos_Ipad2_3to4 = ()   
                break
            else:
                pos_Ipad2_3to4 = random.choice(time_3to4)
        if pos_Ipad2_3to4 in Employee_name:
           Ipad2_3to4 = Employee_name[pos_Ipad2_3to4]
           time_3to4.remove(pos_Ipad2_3to4)
           print((Ipad2_3to4) + " on IPad")
        if time_3to4 == []:
            continue
    #Desserts
        pos_Desserts3to4 = random.choice(time_3to4)
        if pos_Desserts3to4 in Employee_name:
           Desserts3to4 = Employee_name[pos_Desserts3to4]
        time_3to4.remove(pos_Desserts3to4)
        print((Desserts3to4) + " on Desserts")
        if time_3to4 == []:
            continue
    #Register 2
        pos_Register2_3to4 = random.choice(time_3to4)
        if pos_Register2_3to4 in Employee_name:
           Register2_3to4 = Employee_name[pos_Register2_3to4]
        time_3to4.remove(pos_Register2_3to4)
        print((Register2_3to4) + " on Register")
        if time_3to4 == []:
            continue
    #IPad 3
        pos_Ipad3_3to4 = random.choice(time_3to4)
        while pos_Ipad3_3to4 == pos_Ipad1_2to3 or pos_Ipad3_3to4 == pos_Ipad2_2to3 or pos_Ipad3_3to4 == pos_Ipad3_2to3:
                if len(time_3to4) == 1:
                    pos_Ipad3_3to4 = ()   
                    break
                else:
                    pos_Ipad3_3to4 = random.choice(time_3to4)
        if pos_Ipad3_3to4 in Employee_name:
           Ipad3_3to4 = Employee_name[pos_Ipad3_3to4]
           time_3to4.remove(pos_Ipad3_3to4)
           print((Ipad3_3to4) + " on IPad")
        if time_3to4 == []:
            continue
    #Dining Room
        pos_Dining3to4 = random.choice(time_3to4)
        if pos_Dining3to4 in Employee_name:
           Dining3to4 = Employee_name[pos_Dining3to4]
        time_3to4.remove(pos_Dining3to4)
        print((Dining3to4) + " on Dining Room")
        if time_3to4 == []:
            continue
    #Sauces
        pos_Sauce3to4 = random.choice(time_3to4)
        if pos_Sauce3to4 in Employee_name:
           Sauce3to4 = Employee_name[pos_Sauce3to4]
        time_3to4.remove(pos_Sauce3to4)
        print((Sauce3to4) + " on Sauces") 
        if time_3to4 == []:
            break
#4 to 5
print ("\nFor 4 pm to 5 pm:")
for x in time_4to5:
    while time_4to5 != []:
    #Register 1
        if '42' in time_4to5:
           print("Taylor on Register")
           time_4to5.remove('42') 
        else:
            pos_Register1_4to5 = random.choice(time_4to5)
            if pos_Register1_4to5 in Employee_name:
             Register1_4to5 = Employee_name[pos_Register1_4to5]
             time_4to5.remove(pos_Register1_4to5)
             print((Register1_4to5) + " on Register") 
        if time_4to5 == []:
             continue     
    #Window    
        pos_window4to5 = random.choice(time_4to5)
        if pos_window4to5 in Employee_name:
           window4to5 = Employee_name[pos_window4to5]
        time_4to5.remove(pos_window4to5)
        print((window4to5) + " on Window")
        if time_4to5 == []:
            continue
    #DT bagging    
        pos_DTbag4to5 = random.choice(time_4to5)
        if pos_DTbag4to5 in Employee_name:
           DTbag4to5 = Employee_name[pos_DTbag4to5]
        time_4to5.remove(pos_DTbag4to5)
        print((DTbag4to5) + " on DT Bagging")
        if time_4to5 == []:
            continue
    #Drinks
        pos_Drinks4to5 = random.choice(time_4to5)
        if pos_Drinks4to5 in Employee_name:
           Drinks4to5 = Employee_name[pos_Drinks4to5]
        time_4to5.remove(pos_Drinks4to5)
        print((Drinks4to5) + " on Drinks")
        if time_4to5 == []:
            continue
    #IPad 1
        pos_Ipad1_4to5 = ()
        pos_Ipad2_4to5 = ()
        pos_Ipad3_4to5 = ()
        pos_Ipad1_4to5 = random.choice(time_4to5)
        while pos_Ipad1_4to5 == pos_Ipad1_3to4 or pos_Ipad1_4to5 == pos_Ipad2_3to4 or pos_Ipad1_4to5 == pos_Ipad3_3to4:
                if len(time_4to5) == 1:
                    pos_Ipad1_4to5 = ()   
                    break
                else:
                    pos_Ipad1_4to5 = random.choice(time_4to5)
        if pos_Ipad1_4to5 in Employee_name:
           Ipad1_4to5 = Employee_name[pos_Ipad1_4to5]
           time_4to5.remove(pos_Ipad1_4to5)
           print((Ipad1_4to5) + " on IPad")
        if time_4to5 == []:
            continue
    #Cash
        pos_Cash4to5 = random.choice(time_4to5)
        if pos_Cash4to5 in Employee_name:
           Cash4to5 = Employee_name[pos_Cash4to5]
        time_4to5.remove(pos_Cash4to5)
        print((Cash4to5) + " on Cash")
        if time_4to5 == []:
            continue
    #FC bagging
        pos_FCbag4to5 = random.choice(time_4to5)
        if pos_FCbag4to5 in Employee_name:
           FCbag4to5 = Employee_name[pos_FCbag4to5]
        time_4to5.remove(pos_FCbag4to5)
        print((FCbag4to5) + " on FC Bagging")
        if time_4to5 == []:
            continue
    #IPad 2
        pos_Ipad2_4to5 = random.choice(time_4to5)
        while pos_Ipad2_4to5 == pos_Ipad1_3to4 or pos_Ipad2_4to5 == pos_Ipad2_3to4 or pos_Ipad2_4to5 == pos_Ipad3_3to4:   
            if len(time_4to5) == 1:
                pos_Ipad2_4to5 = ()   
                break
            else:
                pos_Ipad2_4to5 = random.choice(time_4to5)
        if pos_Ipad2_4to5 in Employee_name:
           Ipad2_4to5 = Employee_name[pos_Ipad2_4to5]
           time_4to5.remove(pos_Ipad2_4to5)
           print((Ipad2_4to5) + " on IPad")
        if time_4to5 == []:
            continue
    #Desserts
        pos_Desserts4to5 = random.choice(time_4to5)
        if pos_Desserts4to5 in Employee_name:
           Desserts4to5 = Employee_name[pos_Desserts4to5]
        time_4to5.remove(pos_Desserts4to5)
        print((Desserts4to5) + " on Desserts")
        if time_4to5 == []:
            continue
    #Register 2
        pos_Register2_4to5 = random.choice(time_4to5)
        if pos_Register2_4to5 in Employee_name:
           Register2_4to5 = Employee_name[pos_Register2_4to5]
        time_4to5.remove(pos_Register2_4to5)
        print((Register2_4to5) + " on Register")
        if time_4to5 == []:
            continue
    #IPad 3
        pos_Ipad3_4to5 = random.choice(time_4to5)
        while pos_Ipad3_4to5 == pos_Ipad1_3to4 or pos_Ipad3_4to5 == pos_Ipad2_3to4 or pos_Ipad3_4to5 == pos_Ipad3_3to4:  
            if len(time_4to5) == 1:
                pos_Ipad3_4to5 = ()    
                break
            else:
                pos_Ipad3_4to5 = random.choice(time_4to5)
        if pos_Ipad3_4to5 in Employee_name:
           Ipad3_4to5 = Employee_name[pos_Ipad3_4to5]
           time_4to5.remove(pos_Ipad3_4to5)
           print((Ipad3_4to5) + " on IPad")
        if time_4to5 == []:
            continue
    #Dining Room
        pos_Dining4to5 = random.choice(time_4to5)
        if pos_Dining4to5 in Employee_name:
           Dining4to5 = Employee_name[pos_Dining4to5]
        time_4to5.remove(pos_Dining4to5)
        print((Dining4to5) + " on Dining Room")
        if time_4to5 == []:
            continue
    #Sauces
        pos_Sauce4to5 = random.choice(time_4to5)
        if pos_Sauce4to5 in Employee_name:
           Sauce4to5 = Employee_name[pos_Sauce4to5]
        time_4to5.remove(pos_Sauce4to5)
        print((Sauce4to5) + " on Sauces") 
        if time_4to5 == []:
            break
#5 to 6
print ("\nFor 5 pm to 6 pm:")
for x in time_5to6:
    while time_5to6 != []:
    #Register 1
        if '42' in time_5to6:
           print("Taylor on Register")
           time_5to6.remove('42') 
        else:
            pos_Register1_5to6 = random.choice(time_5to6)
            if pos_Register1_5to6 in Employee_name:
             Register1_5to6 = Employee_name[pos_Register1_5to6]
             time_5to6.remove(pos_Register1_5to6)
             print((Register1_5to6) + " on Register") 
        if time_5to6 == []:
             continue     
    #Window    
        pos_window5to6 = random.choice(time_5to6)
        if pos_window5to6 in Employee_name:
           window5to6 = Employee_name[pos_window5to6]
        time_5to6.remove(pos_window5to6)
        print((window5to6) + " on Window")
        if time_5to6 == []:
            continue
    #DT bagging    
        pos_DTbag5to6 = random.choice(time_5to6)
        if pos_DTbag5to6 in Employee_name:
           DTbag5to6 = Employee_name[pos_DTbag5to6]
        time_5to6.remove(pos_DTbag5to6)
        print((DTbag5to6) + " on DT Bagging")
        if time_5to6 == []:
            continue
    #Drinks
        pos_Drinks5to6 = random.choice(time_5to6)
        if pos_Drinks5to6 in Employee_name:
           Drinks5to6 = Employee_name[pos_Drinks5to6]
        time_5to6.remove(pos_Drinks5to6)
        print((Drinks5to6) + " on Drinks")
        if time_5to6 == []:
            continue
    #IPad 1
        pos_Ipad1_5to6 = ()
        pos_Ipad2_5to6 = ()
        pos_Ipad3_5to6 = ()
        pos_Ipad1_5to6 = random.choice(time_5to6)
        while pos_Ipad1_5to6 == pos_Ipad1_4to5 or pos_Ipad1_5to6 == pos_Ipad2_4to5 or pos_Ipad1_5to6 == pos_Ipad3_4to5:
            if len(time_5to6) == 1:
                pos_Ipad1_5to6 = ()   
                break
            else:
                pos_Ipad1_5to6 = random.choice(time_5to6)
        if pos_Ipad1_5to6 in Employee_name:
           Ipad1_5to6 = Employee_name[pos_Ipad1_5to6]
           time_5to6.remove(pos_Ipad1_5to6)
           print((Ipad1_5to6) + " on IPad")
        if time_5to6 == []:
            continue
    #Cash
        pos_Cash5to6 = random.choice(time_5to6)
        if pos_Cash5to6 in Employee_name:
           Cash5to6 = Employee_name[pos_Cash5to6]
        time_5to6.remove(pos_Cash5to6)
        print((Cash5to6) + " on Cash")
        if time_5to6 == []:
            continue
    #FC bagging
        pos_FCbag5to6 = random.choice(time_5to6)
        if pos_FCbag5to6 in Employee_name:
           FCbag5to6 = Employee_name[pos_FCbag5to6]
        time_5to6.remove(pos_FCbag5to6)
        print((FCbag5to6) + " on FC Bagging")
        if time_5to6 == []:
            continue
    #IPad 2
        pos_Ipad2_5to6 = random.choice(time_5to6)
        while pos_Ipad2_5to6 == pos_Ipad1_4to5 or pos_Ipad2_5to6 == pos_Ipad2_4to5 or pos_Ipad2_5to6 == pos_Ipad3_4to5:   
            if len(time_5to6) == 1:
                pos_Ipad2_5to6 = ()   
                break
            else:
                pos_Ipad2_5to6 = random.choice(time_5to6)
        if pos_Ipad2_5to6 in Employee_name:
           Ipad2_5to6 = Employee_name[pos_Ipad2_5to6]
           time_5to6.remove(pos_Ipad2_5to6)
           print((Ipad2_5to6) + " on IPad")
        if time_5to6 == []:
            continue
    #Desserts
        pos_Desserts5to6 = random.choice(time_5to6)
        if pos_Desserts5to6 in Employee_name:
           Desserts5to6 = Employee_name[pos_Desserts5to6]
        time_5to6.remove(pos_Desserts5to6)
        print((Desserts5to6) + " on Desserts")
        if time_5to6 == []:
            continue
    #Register 2
        pos_Register2_5to6 = random.choice(time_5to6)
        if pos_Register2_5to6 in Employee_name:
           Register2_5to6 = Employee_name[pos_Register2_5to6]
        time_5to6.remove(pos_Register2_5to6)
        print((Register2_5to6) + " on Register")
        if time_5to6 == []:
            continue
    #IPad 3
        pos_Ipad3_5to6 = random.choice(time_5to6)
        while pos_Ipad3_5to6 == pos_Ipad1_4to5 or pos_Ipad3_5to6 == pos_Ipad2_4to5 or pos_Ipad3_5to6 == pos_Ipad3_4to5:
                if len(time_5to6) == 1:
                    pos_Ipad3_5to6 = ()   
                    break
                else:
                    pos_Ipad3_5to6 = random.choice(time_5to6)
        if pos_Ipad3_5to6 in Employee_name:
           Ipad3_5to6 = Employee_name[pos_Ipad3_5to6]
           time_5to6.remove(pos_Ipad3_5to6)
           print((Ipad3_5to6) + " on IPad")
        if time_5to6 == []:
            continue
    #Dining Room
        pos_Dining5to6 = random.choice(time_5to6)
        if pos_Dining5to6 in Employee_name:
           Dining5to6 = Employee_name[pos_Dining5to6]
        time_5to6.remove(pos_Dining5to6)
        print((Dining5to6) + " on Dining Room")
        if time_5to6 == []:
            continue
    #Sauces
        pos_Sauce5to6 = random.choice(time_5to6)
        if pos_Sauce5to6 in Employee_name:
           Sauce5to6 = Employee_name[pos_Sauce5to6]
        time_5to6.remove(pos_Sauce5to6)
        print((Sauce5to6) + " on Sauces") 
        if time_5to6 == []:
            break
#6 to 7
print ("\nFor 6 pm to 7 pm:")              
for x in time_6to7:
    while time_6to7 != []:
    #Register 1
        if '42' in time_6to7:
           print("Taylor on Register")
           time_6to7.remove('42') 
        else:
            pos_Register1_6to7 = random.choice(time_6to7)
            if pos_Register1_6to7 in Employee_name:
             Register1_6to7 = Employee_name[pos_Register1_6to7]
             time_6to7.remove(pos_Register1_6to7)
             print((Register1_6to7) + " on Register") 
        if time_6to7 == []:
             continue     
#Window    
        pos_window6to7 = random.choice(time_6to7)
        if pos_window6to7 in Employee_name:
           window6to7 = Employee_name[pos_window6to7]
        time_6to7.remove(pos_window6to7)
        print((window6to7) + " on Window")
        if time_6to7 == []:
            continue
    #DT bagging    
        pos_DTbag6to7 = random.choice(time_6to7)
        if pos_DTbag6to7 in Employee_name:
           DTbag6to7 = Employee_name[pos_DTbag6to7]
        time_6to7.remove(pos_DTbag6to7)
        print((DTbag6to7) + " on DT Bagging")
        if time_6to7 == []:
            continue
    #Drinks
        pos_Drinks6to7 = random.choice(time_6to7)
        if pos_Drinks6to7 in Employee_name:
           Drinks6to7 = Employee_name[pos_Drinks6to7]
        time_6to7.remove(pos_Drinks6to7)
        print((Drinks6to7) + " on Drinks")
        if time_6to7 == []:
            continue
    #IPad 1
        pos_Ipad1_6to7 = ()
        pos_Ipad2_6to7 = ()
        pos_Ipad3_6to7 = ()
        pos_Ipad1_6to7 = random.choice(time_6to7)
        while pos_Ipad1_6to7 == pos_Ipad1_5to6 or pos_Ipad1_6to7 == pos_Ipad2_5to6 or pos_Ipad1_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad1_6to7 = ()   
                break
            else:
                pos_Ipad1_6to7 = random.choice(time_6to7)           
        if pos_Ipad1_6to7 in Employee_name:
           Ipad1_6to7 = Employee_name[pos_Ipad1_6to7]
           time_6to7.remove(pos_Ipad1_6to7)
           print((Ipad1_6to7) + " on IPad")
        if time_6to7 == []:
            continue
    #Cash
        pos_Cash6to7 = random.choice(time_6to7)
        if pos_Cash6to7 in Employee_name:
           Cash6to7 = Employee_name[pos_Cash6to7]
        time_6to7.remove(pos_Cash6to7)
        print((Cash6to7) + " on Cash")
        if time_6to7 == []:
            continue
    #FC bagging
        pos_FCbag6to7 = random.choice(time_6to7)
        if pos_FCbag6to7 in Employee_name:
           FCbag6to7 = Employee_name[pos_FCbag6to7]
        time_6to7.remove(pos_FCbag6to7)
        print((FCbag6to7) + " on FC Bagging")
        if time_6to7 == []:
            continue
    #IPad 2
        pos_Ipad2_6to7 = random.choice(time_6to7)
        while pos_Ipad2_6to7 == pos_Ipad1_5to6 or pos_Ipad2_6to7 == pos_Ipad2_5to6 or pos_Ipad2_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad2_6to7 = ()   
                break
            else:
                pos_Ipad2_6to7 = random.choice(time_6to7)   
        if pos_Ipad2_6to7 in Employee_name:
           Ipad2_6to7 = Employee_name[pos_Ipad2_6to7]
           time_6to7.remove(pos_Ipad2_6to7)
           print((Ipad2_6to7) + " on IPad")
        if time_6to7 == []:
            continue
    #Desserts
        pos_Desserts6to7 = random.choice(time_6to7)
        if pos_Desserts6to7 in Employee_name:
           Desserts6to7 = Employee_name[pos_Desserts6to7]
        time_6to7.remove(pos_Desserts6to7)
        print((Desserts6to7) + " on Desserts")
        if time_6to7 == []:
            continue
    #Register 2
        pos_Register2_6to7 = random.choice(time_6to7)
        if pos_Register2_6to7 in Employee_name:
           Register2_6to7 = Employee_name[pos_Register2_6to7]
        time_6to7.remove(pos_Register2_6to7)
        print((Register2_6to7) + " on Register")
        if time_6to7 == []:
            continue
    #IPad 3
        pos_Ipad3_6to7 = random.choice(time_6to7)
        while pos_Ipad3_6to7 == pos_Ipad1_5to6 or pos_Ipad3_6to7 == pos_Ipad2_5to6 or pos_Ipad3_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad3_6to7 = ()   
                break   
            else:
                pos_Ipad3_6to7 = random.choice(time_6to7)     
        if pos_Ipad3_6to7 in Employee_name:
           Ipad3_6to7 = Employee_name[pos_Ipad3_6to7]
           time_6to7.remove(pos_Ipad3_6to7)
           print((Ipad3_6to7) + " on IPad")
        if time_6to7 == []:
            continue
    #Dining Room
        pos_Dining6to7 = random.choice(time_6to7)
        if pos_Dining6to7 in Employee_name:
           Dining6to7 = Employee_name[pos_Dining6to7]
        time_6to7.remove(pos_Dining6to7)
        print((Dining6to7) + " on Dining Room")
        if time_6to7 == []:
            continue
    #Sauces
        pos_Sauce6to7 = random.choice(time_6to7)
        if pos_Sauce6to7 in Employee_name:
           Sauce6to7 = Employee_name[pos_Sauce6to7]
        time_6to7.remove(pos_Sauce6to7)
        print((Sauce6to7) + " on Sauces") 
        if time_6to7 == []:
            break
        
#7 to 8
print ("\nFor 7 pm to 8 pm:")              
for x in time_7to8:
    while time_7to8 != []:
#Window    
        pos_window7to8 = random.choice(time_7to8)
        if pos_window7to8 in Employee_name:
           window7to8 = Employee_name[pos_window7to8]
        time_7to8.remove(pos_window7to8)
        print((window7to8) + " on Window")
        if time_7to8 == []:
            continue
    #DT bagging    
        pos_DTbag7to8 = random.choice(time_7to8)
        if pos_DTbag7to8 in Employee_name:
           DTbag7to8 = Employee_name[pos_DTbag7to8]
        time_7to8.remove(pos_DTbag7to8)
        print((DTbag7to8) + " on DT Bagging")
        if time_7to8 == []:
            continue
    #Drinks
        pos_Drinks7to8 = random.choice(time_7to8)
        if pos_Drinks7to8 in Employee_name:
           Drinks7to8 = Employee_name[pos_Drinks7to8]
        time_7to8.remove(pos_Drinks7to8)
        print((Drinks7to8) + " on Drinks")
        if time_7to8 == []:
            continue
    #IPad 1
        pos_Ipad1_7to8 = ()
        pos_Ipad2_7to8 = ()
        pos_Ipad3_7to8 = ()
        pos_Ipad1_7to8 = random.choice(time_7to8)
        while pos_Ipad1_7to8 == pos_Ipad1_6to7 or pos_Ipad1_7to8 == pos_Ipad2_6to7 or pos_Ipad1_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad1_7to8 = ()   
                break   
            else:
                pos_Ipad1_7to8 = random.choice(time_7to8)
        if pos_Ipad1_7to8 in Employee_name:
           Ipad1_7to8 = Employee_name[pos_Ipad1_7to8]
           time_7to8.remove(pos_Ipad1_7to8)
           print((Ipad1_7to8) + " on IPad")
        if time_7to8 == []:
            continue
    #Cash
        pos_Cash7to8 = random.choice(time_7to8)
        if pos_Cash7to8 in Employee_name:
           Cash7to8 = Employee_name[pos_Cash7to8]
        time_7to8.remove(pos_Cash7to8)
        print((Cash7to8) + " on Cash")
        if time_7to8 == []:
            continue
    #Register 1
        pos_Register1_7to8 = random.choice(time_7to8)
        if pos_Register1_7to8 in Employee_name:
            Register1_7to8 = Employee_name[pos_Register1_7to8]
            time_7to8.remove(pos_Register1_7to8)
            print((Register1_7to8) + " on Register")
        if time_7to8 == []:
             continue
    #FC bagging
        pos_FCbag7to8 = random.choice(time_7to8)
        if pos_FCbag7to8 in Employee_name:
           FCbag7to8 = Employee_name[pos_FCbag7to8]
        time_7to8.remove(pos_FCbag7to8)
        print((FCbag7to8) + " on FC Bagging")
        if time_7to8 == []:
            continue
    #IPad 2
        pos_Ipad2_7to8 = random.choice(time_7to8)
        while pos_Ipad2_7to8 == pos_Ipad1_6to7 or pos_Ipad2_7to8 == pos_Ipad2_6to7 or pos_Ipad2_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad2_7to8 = ()   
                break   
            else:
                pos_Ipad2_7to8 = random.choice(time_7to8)
        if pos_Ipad2_7to8 in Employee_name:
           Ipad2_7to8 = Employee_name[pos_Ipad2_7to8]
           time_7to8.remove(pos_Ipad2_7to8)
           print((Ipad2_7to8) + " on IPad")
        if time_7to8 == []:
            continue
    #Desserts
        pos_Desserts7to8 = random.choice(time_7to8)
        if pos_Desserts7to8 in Employee_name:
           Desserts7to8 = Employee_name[pos_Desserts7to8]
        time_7to8.remove(pos_Desserts7to8)
        print((Desserts7to8) + " on Desserts")
        if time_7to8 == []:
            continue
    #Register 2
        pos_Register2_7to8 = random.choice(time_7to8)
        if pos_Register2_7to8 in Employee_name:
           Register2_7to8 = Employee_name[pos_Register2_7to8]
        time_7to8.remove(pos_Register2_7to8)
        print((Register2_7to8) + " on Register")
        if time_7to8 == []:
            continue
    #IPad 3
        pos_Ipad3_7to8 = random.choice(time_7to8)
        while pos_Ipad3_7to8 == pos_Ipad1_6to7 or pos_Ipad3_7to8 == pos_Ipad2_6to7 or pos_Ipad3_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad3_7to8 = ()   
                break   
            else:
                pos_Ipad3_7to8 = random.choice(time_7to8)
        if pos_Ipad3_7to8 in Employee_name:
           Ipad3_7to8 = Employee_name[pos_Ipad3_7to8]
           time_7to8.remove(pos_Ipad3_7to8)
           print((Ipad3_7to8) + " on IPad")
        if time_7to8 == []:
            continue
    #Dining Room
        pos_Dining7to8 = random.choice(time_7to8)
        if pos_Dining7to8 in Employee_name:
           Dining7to8 = Employee_name[pos_Dining7to8]
        time_7to8.remove(pos_Dining7to8)
        print((Dining7to8) + " on Dining Room")
        if time_7to8 == []:
            continue
    #Sauces
        pos_Sauce7to8 = random.choice(time_7to8)
        if pos_Sauce7to8 in Employee_name:
           Sauce7to8 = Employee_name[pos_Sauce7to8]
        time_7to8.remove(pos_Sauce7to8)
        print((Sauce7to8) + " on Sauces") 
        if time_7to8 == []:
            break
#8 to 9
print ("\nFor 8 pm to 9 pm:")              
for x in time_8to9:
    while time_8to9 != []:
#Window    
        pos_window8to9 = random.choice(time_8to9)
        if pos_window8to9 in Employee_name:
           window8to9 = Employee_name[pos_window8to9]
        time_8to9.remove(pos_window8to9)
        print((window8to9) + " on Window")
        if time_8to9 == []:
            continue
    #DT bagging    
        pos_DTbag8to9 = random.choice(time_8to9)
        if pos_DTbag8to9 in Employee_name:
           DTbag8to9 = Employee_name[pos_DTbag8to9]
        time_8to9.remove(pos_DTbag8to9)
        print((DTbag8to9) + " on DT Bagging")
        if time_8to9 == []:
            continue
    #Drinks
        pos_Drinks8to9 = random.choice(time_8to9)
        if pos_Drinks8to9 in Employee_name:
           Drinks8to9 = Employee_name[pos_Drinks8to9]
        time_8to9.remove(pos_Drinks8to9)
        print((Drinks8to9) + " on Drinks")
        if time_8to9 == []:
            continue
    #Dining Room
        pos_Dining8to9 = random.choice(time_8to9)
        if pos_Dining8to9 in Employee_name:
           Dining8to9 = Employee_name[pos_Dining8to9]
        time_8to9.remove(pos_Dining8to9)
        print((Dining8to9) + " on Dining Room")
        if time_8to9 == []:
            continue        
    #IPad 1
        pos_Ipad1_8to9 = ()
        pos_Ipad2_8to9 = ()
        pos_Ipad3_8to9 = ()
        pos_Ipad1_8to9 = random.choice(time_8to9)
        while pos_Ipad1_8to9 == pos_Ipad1_7to8 or pos_Ipad1_8to9 == pos_Ipad2_7to8 or pos_Ipad1_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad1_8to9 = ()   
                break  
            else:
                pos_Ipad1_8to9 = random.choice(time_8to9)
        if pos_Ipad1_8to9 in Employee_name:
           Ipad1_8to9 = Employee_name[pos_Ipad1_8to9]
           time_8to9.remove(pos_Ipad1_8to9)
           print((Ipad1_8to9) + " on IPad")
        if time_8to9 == []:
            continue
    #Cash
        pos_Cash8to9 = random.choice(time_8to9)
        if pos_Cash8to9 in Employee_name:
           Cash8to9 = Employee_name[pos_Cash8to9]
        time_8to9.remove(pos_Cash8to9)
        print((Cash8to9) + " on Cash")
        if time_8to9 == []:
            continue
    #Register 1
        pos_Register1_8to9 = random.choice(time_8to9)
        if pos_Register1_8to9 in Employee_name:
            Register1_8to9 = Employee_name[pos_Register1_8to9]
            time_8to9.remove(pos_Register1_8to9)
            print((Register1_8to9) + " on Register")
        if time_8to9 == []:
             continue
    #FC bagging
        pos_FCbag8to9 = random.choice(time_8to9)
        if pos_FCbag8to9 in Employee_name:
           FCbag8to9 = Employee_name[pos_FCbag8to9]
        time_8to9.remove(pos_FCbag8to9)
        print((FCbag8to9) + " on FC Bagging")
        if time_8to9 == []:
            continue
    #IPad 2
        pos_Ipad2_8to9 = random.choice(time_8to9)
        while pos_Ipad2_8to9 == pos_Ipad1_7to8 or pos_Ipad2_8to9 == pos_Ipad2_7to8 or pos_Ipad2_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad2_8to9 = ()   
                break  
            else:
                pos_Ipad2_8to9 = random.choice(time_8to9)
        if pos_Ipad2_8to9 in Employee_name:
           Ipad2_8to9 = Employee_name[pos_Ipad2_8to9]
           time_8to9.remove(pos_Ipad2_8to9)
           print((Ipad2_8to9) + " on IPad")
        if time_8to9 == []:
            continue
    #Desserts
        pos_Desserts8to9 = random.choice(time_8to9)
        if pos_Desserts8to9 in Employee_name:
           Desserts8to9 = Employee_name[pos_Desserts8to9]
        time_8to9.remove(pos_Desserts8to9)
        print((Desserts8to9) + " on Desserts")
        if time_8to9 == []:
            continue
    #Register 2
        pos_Register2_8to9 = random.choice(time_8to9)
        if pos_Register2_8to9 in Employee_name:
           Register2_8to9 = Employee_name[pos_Register2_8to9]
        time_8to9.remove(pos_Register2_8to9)
        print((Register2_8to9) + " on Register")
        if time_8to9 == []:
            continue
    #IPad 3
        pos_Ipad3_8to9 = random.choice(time_8to9)
        while pos_Ipad3_8to9 == pos_Ipad1_7to8 or pos_Ipad3_8to9 == pos_Ipad2_7to8 or pos_Ipad3_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad3_8to9 = ()   
                break  
            else:
                pos_Ipad3_8to9 = random.choice(time_8to9)
        if pos_Ipad3_8to9 in Employee_name:
           Ipad3_8to9 = Employee_name[pos_Ipad3_8to9]
           time_8to9.remove(pos_Ipad3_8to9)
           print((Ipad3_8to9) + " on IPad")
        if time_8to9 == []:
            continue
    #Sauces
        pos_Sauce8to9 = random.choice(time_8to9)
        if pos_Sauce8to9 in Employee_name:
           Sauce8to9 = Employee_name[pos_Sauce8to9]
        time_8to9.remove(pos_Sauce8to9)
        print((Sauce8to9) + " on Sauces") 
        if time_8to9 == []:
            break        
#9 to 10
print ("\nFor 9 pm to 10 pm:")              
for x in time_9to10:
    while time_9to10 != []:
     #Window    
        pos_window9to10 = random.choice(time_9to10)
        if pos_window9to10 in Employee_name:
           window9to10 = Employee_name[pos_window9to10]
        time_9to10.remove(pos_window9to10)
        print((window9to10) + " on Window")
        if time_9to10 == []:
            continue
    #DT bagging    
        pos_DTbag9to10 = random.choice(time_9to10)
        if pos_DTbag9to10 in Employee_name:
           DTbag9to10 = Employee_name[pos_DTbag9to10]
        time_9to10.remove(pos_DTbag9to10)
        print((DTbag9to10) + " on DT Bagging")
        if time_9to10 == []:
            continue
    #Drinks
        pos_Drinks9to10 = random.choice(time_9to10)
        if pos_Drinks9to10 in Employee_name:
           Drinks9to10 = Employee_name[pos_Drinks9to10]
        time_9to10.remove(pos_Drinks9to10)
        print((Drinks9to10) + " on Drinks")
        if time_9to10 == []:
            continue
    #Dining Room
        pos_Dining9to10 = random.choice(time_9to10)
        if pos_Dining9to10 in Employee_name:
           Dining9to10 = Employee_name[pos_Dining9to10]
        time_9to10.remove(pos_Dining9to10)
        print((Dining9to10) + " on Dining Room")
        if time_9to10 == []:
            continue
    #Register 1
        pos_Register1_9to10 = random.choice(time_9to10)
        if pos_Register1_9to10 in Employee_name:
            Register1_9to10 = Employee_name[pos_Register1_9to10]
            time_9to10.remove(pos_Register1_9to10)
            print((Register1_9to10) + " on Register")
        if time_9to10 == []:
             continue
    #FC bagging
        pos_FCbag9to10 = random.choice(time_9to10)
        if pos_FCbag9to10 in Employee_name:
           FCbag9to10 = Employee_name[pos_FCbag9to10]
        time_9to10.remove(pos_FCbag9to10)
        print((FCbag9to10) + " on FC Bagging")
        if time_9to10 == []:
            continue

#Output as word document

def iter_paragraphs_of_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                  yield paragraph
                yield from iter_paragraphs_of_tables(cell.tables)

for paragraph in iter_paragraphs_of_tables(doc.tables):
     if '{{Win2to3}}' in paragraph.text:
      if window2to3 != ():
       paragraph.text = paragraph.text.replace("{{Win2to3}}", window2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Win2to3}}", "")   
     if '{{DTBag2to3}}' in paragraph.text:
      if DTbag2to3 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", DTbag2to3 + " 2-3")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", "")
     if '{{Drinks2to3}}' in paragraph.text:
      if Drinks2to3 != ():
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", Drinks2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", "")   
     if '{{Desserts2to3}}' in paragraph.text:
        if Desserts2to3 != ():
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", Desserts2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", "")   
     if '{{Sauce2to3}}' in paragraph.text:
        if Sauce2to3 != ():
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", Sauce2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", "")   
     if '{{FCBag2to3}}' in paragraph.text:
        if FCbag2to3 != ():
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", FCbag2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", "")   
     if '{{Reg1_2to3}}' in paragraph.text:
        if Register1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", Register1_2to3 + " 2-3")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", "") 
     if '{{Reg2_2to3}}' in paragraph.text:
        if Register2_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", Register2_2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", "")  
     if '{{Dining_2to3}}' in paragraph.text:
        if Dining2to3 != ():
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", Dining2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", "")
     if '{{Cash2to3}}' in paragraph.text:
        if Cash2to3 != ():
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", Cash2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", "")   
     if '{{IPad1_2to3}}' in paragraph.text:
        if Ipad1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", Ipad1_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", "")    
     if '{{IPad2_2to3}}' in paragraph.text:
        if Ipad2_2to3 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", Ipad2_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", "")   
     if '{{IPad3_2to3}}' in paragraph.text:
        if Ipad3_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", Ipad3_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", "")
     if '{{Win3to4}}' in paragraph.text:
      if window3to4 != ():
       paragraph.text = paragraph.text.replace("{{Win3to4}}", window3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Win3to4}}", "")   
     if '{{DTBag3to4}}' in paragraph.text:
      if DTbag3to4 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", DTbag3to4 + " 3-4")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", "")
     if '{{Drinks3to4}}' in paragraph.text:
      if Drinks3to4 != ():
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", Drinks3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", "")   
     if '{{Desserts3to4}}' in paragraph.text:
        if Desserts3to4 != ():
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", Desserts3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", "")   
     if '{{Sauce3to4}}' in paragraph.text:
        if Sauce3to4 != ():
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", Sauce3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", "")   
     if '{{FCBag3to4}}' in paragraph.text:
        if FCbag3to4 != ():
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", FCbag3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", "")   
     if '{{Reg1_3to4}}' in paragraph.text:
        if Register1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", Register1_3to4 + " 3-4")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", "") 
     if '{{Reg2_3to4}}' in paragraph.text:
        if Register2_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", Register2_3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", "")  
     if '{{Dining_3to4}}' in paragraph.text:
        if Dining3to4 != ():
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", Dining3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", "")
     if '{{Cash3to4}}' in paragraph.text:
        if Cash3to4 != ():
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", Cash3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", "")   
     if '{{IPad1_3to4}}' in paragraph.text:
        if Ipad1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", Ipad1_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", "")    
     if '{{IPad2_3to4}}' in paragraph.text:
        if Ipad2_3to4 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", Ipad2_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", "")   
     if '{{IPad3_3to4}}' in paragraph.text:
        if Ipad3_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", Ipad3_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", "")       
     if '{{Win4to5}}' in paragraph.text:
      if window4to5 != ():
       paragraph.text = paragraph.text.replace("{{Win4to5}}", window4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Win4to5}}", "")   
     if '{{DTBag4to5}}' in paragraph.text:
      if DTbag4to5 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", DTbag4to5 + " 4-5")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", "")
     if '{{Drinks4to5}}' in paragraph.text:
      if Drinks4to5 != ():
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", Drinks4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", "")   
     if '{{Desserts4to5}}' in paragraph.text:
        if Desserts4to5 != ():
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", Desserts4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", "")   
     if '{{Sauce4to5}}' in paragraph.text:
        if Sauce4to5 != ():
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", Sauce4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", "")   
     if '{{FCBag4to5}}' in paragraph.text:
        if FCbag4to5 != ():
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", FCbag4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", "")   
     if '{{Reg1_4to5}}' in paragraph.text:
        if Register1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", Register1_4to5 + " 4-5")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", "") 
     if '{{Reg2_4to5}}' in paragraph.text:
        if Register2_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", Register2_4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", "")  
     if '{{Dining_4to5}}' in paragraph.text:
        if Dining4to5 != ():
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", Dining4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", "")
     if '{{Cash4to5}}' in paragraph.text:
        if Cash4to5 != ():
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", Cash4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", "")   
     if '{{IPad1_4to5}}' in paragraph.text:
        if Ipad1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", Ipad1_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", "")    
     if '{{IPad2_4to5}}' in paragraph.text:
        if Ipad2_4to5 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", Ipad2_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", "")   
     if '{{IPad3_4to5}}' in paragraph.text:
        if Ipad3_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", Ipad3_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", "")
     if '{{Win5to6}}' in paragraph.text:
      if window5to6 != ():
       paragraph.text = paragraph.text.replace("{{Win5to6}}", window5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Win5to6}}", "")   
     if '{{DTBag5to6}}' in paragraph.text:
      if DTbag5to6 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", DTbag5to6 + " 5-6")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", "")
     if '{{Drinks5to6}}' in paragraph.text:
      if Drinks5to6 != ():
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", Drinks5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", "")   
     if '{{Desserts5to6}}' in paragraph.text:
        if Desserts5to6 != ():
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", Desserts5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", "")   
     if '{{Sauce5to6}}' in paragraph.text:
        if Sauce5to6 != ():
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", Sauce5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", "")   
     if '{{FCBag5to6}}' in paragraph.text:
        if FCbag5to6 != ():
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", FCbag5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", "")   
     if '{{Reg1_5to6}}' in paragraph.text:
        if Register1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", Register1_5to6 + " 5-6")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", "") 
     if '{{Reg2_5to6}}' in paragraph.text:
        if Register2_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", Register2_5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", "")  
     if '{{Dining_5to6}}' in paragraph.text:
        if Dining5to6 != ():
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", Dining5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", "")
     if '{{Cash5to6}}' in paragraph.text:
        if Cash5to6 != ():
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", Cash5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", "")   
     if '{{IPad1_5to6}}' in paragraph.text:
        if Ipad1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", Ipad1_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", "")    
     if '{{IPad2_5to6}}' in paragraph.text:
        if Ipad2_5to6 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", Ipad2_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", "")   
     if '{{IPad3_5to6}}' in paragraph.text:
        if Ipad3_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", Ipad3_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", "")         
     if '{{Win6to7}}' in paragraph.text:
      if window6to7 != ():
       paragraph.text = paragraph.text.replace("{{Win6to7}}", window6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Win6to7}}", "")   
     if '{{DTBag6to7}}' in paragraph.text:
      if DTbag6to7 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", DTbag6to7 + " 6-7")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", "")
     if '{{Drinks6to7}}' in paragraph.text:
      if Drinks6to7 != ():
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", Drinks6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", "")   
     if '{{Desserts6to7}}' in paragraph.text:
        if Desserts6to7 != ():
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", Desserts6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", "")   
     if '{{Sauce6to7}}' in paragraph.text:
        if Sauce6to7 != ():
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", Sauce6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", "")   
     if '{{FCBag6to7}}' in paragraph.text:
        if FCbag6to7 != ():
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", FCbag6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", "")   
     if '{{Reg1_6to7}}' in paragraph.text:
        if Register1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", Register1_6to7 + " 6-7")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", "") 
     if '{{Reg2_6to7}}' in paragraph.text:
        if Register2_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", Register2_6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", "")  
     if '{{Dining_6to7}}' in paragraph.text:
        if Dining6to7 != ():
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", Dining6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", "")
     if '{{Cash6to7}}' in paragraph.text:
        if Cash6to7 != ():
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", Cash6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", "")   
     if '{{IPad1_6to7}}' in paragraph.text:
        if Ipad1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", Ipad1_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", "")    
     if '{{IPad2_6to7}}' in paragraph.text:
        if Ipad2_6to7 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", Ipad2_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", "")   
     if '{{IPad3_6to7}}' in paragraph.text:
        if Ipad3_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", Ipad3_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", "")    
     if '{{Win7to8}}' in paragraph.text:
      if window7to8 != ():
       paragraph.text = paragraph.text.replace("{{Win7to8}}", window7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Win7to8}}", "")   
     if '{{DTBag7to8}}' in paragraph.text:
      if DTbag7to8 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", DTbag7to8 + " 7-8")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", "")
     if '{{Drinks7to8}}' in paragraph.text:
      if Drinks7to8 != ():
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", Drinks7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", "")   
     if '{{Desserts7to8}}' in paragraph.text:
        if Desserts7to8 != ():
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", Desserts7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", "")   
     if '{{Sauce7to8}}' in paragraph.text:
        if Sauce7to8 != ():
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", Sauce7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", "")   
     if '{{FCBag7to8}}' in paragraph.text:
        if FCbag7to8 != ():
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", FCbag7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", "")   
     if '{{Reg1_7to8}}' in paragraph.text:
        if Register1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", Register1_7to8 + " 7-8")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", "") 
     if '{{Reg2_7to8}}' in paragraph.text:
        if Register2_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", Register2_7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", "")  
     if '{{Dining_7to8}}' in paragraph.text:
        if Dining7to8 != ():
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", Dining7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", "")
     if '{{Cash7to8}}' in paragraph.text:
        if Cash7to8 != ():
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", Cash7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", "")   
     if '{{IPad1_7to8}}' in paragraph.text:
        if Ipad1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", Ipad1_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", "")    
     if '{{IPad2_7to8}}' in paragraph.text:
        if Ipad2_7to8 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", Ipad2_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", "")   
     if '{{IPad3_7to8}}' in paragraph.text:
        if Ipad3_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", Ipad3_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", "")
     if '{{Win8to9}}' in paragraph.text:
      if window8to9 != ():
       paragraph.text = paragraph.text.replace("{{Win8to9}}", window8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Win8to9}}", "")   
     if '{{DTBag8to9}}' in paragraph.text:
      if DTbag8to9 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", DTbag8to9 + " 8-9")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", "")
     if '{{Drinks8to9}}' in paragraph.text:
      if Drinks8to9 != ():
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", Drinks8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", "")   
     if '{{Desserts8to9}}' in paragraph.text:
        if Desserts8to9 != ():
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", Desserts8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", "")   
     if '{{Sauce8to9}}' in paragraph.text:
        if Sauce8to9 != ():
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", Sauce8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", "")   
     if '{{FCBag8to9}}' in paragraph.text:
        if FCbag8to9 != ():
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", FCbag8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", "")   
     if '{{Reg1_8to9}}' in paragraph.text:
        if Register1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", Register1_8to9 + " 8-9")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", "") 
     if '{{Reg2_8to9}}' in paragraph.text:
        if Register2_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", Register2_8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", "")  
     if '{{Dining_8to9}}' in paragraph.text:
        if Dining8to9 != ():
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", Dining8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", "")
     if '{{Cash8to9}}' in paragraph.text:
        if Cash8to9 != ():
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", Cash8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", "")   
     if '{{IPad1_8to9}}' in paragraph.text:
        if Ipad1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", Ipad1_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", "")    
     if '{{IPad2_8to9}}' in paragraph.text:
        if Ipad2_8to9 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", Ipad2_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", "")   
     if '{{IPad3_8to9}}' in paragraph.text:
        if Ipad3_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", Ipad3_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", "")         
     if '{{Win9to10}}' in paragraph.text:
      if window9to10 != ():
       paragraph.text = paragraph.text.replace("{{Win9to10}}", window9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Win9to10}}", "")   
     if '{{DTBag9to10}}' in paragraph.text:
      if DTbag9to10 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", DTbag9to10 + " 9-10")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", "")
     if '{{Drinks9to10}}' in paragraph.text:
      if Drinks9to10 != ():
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", Drinks9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", "")   
     if '{{Desserts9to10}}' in paragraph.text:
        if Desserts9to10 != ():
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", Desserts9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", "")   
     if '{{Sauce9to10}}' in paragraph.text:
        if Sauce9to10 != ():
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", Sauce9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", "")   
     if '{{FCBag9to10}}' in paragraph.text:
        if FCbag9to10 != ():
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", FCbag9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", "")   
     if '{{Reg1_9to10}}' in paragraph.text:
        if Register1_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", Register1_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", "") 
     if '{{Reg2_9to10}}' in paragraph.text:
        if Register2_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", Register2_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", "")  
     if '{{Dining_9to10}}' in paragraph.text:
        if Dining9to10 != ():
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", Dining9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", "")
     for y in range (0,len (FOHCrew)):
         team_order = '{{FOHCrew'+ str (y) + '}}'
         if team_order in paragraph.text:
             paragraph.text = paragraph.text.replace(team_order, FOHCrew[y])
if '{{Date}}' in paragraph.text:
  paragraph.text = paragraph.text.replace("{{Date}}", print_date)            
font = paragraph.style.font         
font.size = Pt(9)         
doc.save('ShiftPlanner.docx')         
os.startfile('C:/Users/Brand/Documents/ShiftPlanner.docx')