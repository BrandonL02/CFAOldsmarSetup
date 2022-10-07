import os

import datetime as dt

import calendar

import random

import pandas as pd

import time

from selenium import webdriver

from selenium.webdriver.common.by import By

#Open document for editing
from docx import Document
from docx.shared import Pt 

doc = Document('/Users/admin/Desktop/ShiftPlanner.docx')

tables = doc.tables 
#Get date and weekday within python
date = dt.date.today()

print_date = (calendar.day_name[date.weekday()] + ', ' + date.strftime("%B %d"))

Workers = []
Begin = []
End = []
Duration = []


def minute_offset(user_var):
    if user_var >= 30:
        user_var = 1
    else:
        user_var = 0    
    return int(user_var)    

def time_to_list(Enter):
    Hours_int = []
    for e in Enter:
        e = time.strptime(e,'%H:%M%p')
        if 0 in e[:5]:
            e = e[3]
        else:
           estimate = minute_offset(e[4])
           e = e[3] + estimate
           if e > 12:
               e = e - 12
        Hours_int.append(e)
    return Hours_int
    
    

def Separate(info):
    for key, value in dict(info).items():
        Workers.append(key)
        Duration.append(value)
    for x in Duration:
        on, off = x.split(' - ', 2)
        Begin.append(on)
        End.append(off)
    x = time_to_list(Begin)
    y = time_to_list(End)
    z = Workers
    return x,y,z
    
def remove_unscheduled(Day):
    for key, value in dict(Day).items():
        if value == '':
            del Day[key] 
            
def DaySort(Weekday):
    for h in Weekday:
        Working = (Weekday[h])
        remove_unscheduled(Working)
        x,y,z = Separate(Working)
    return x,y,z
     

#Access Hotschedules with credentials and navigate pages

username = '2445922'

password = 'Monkey2002'

url = "https://www.hotschedules.com/hs/login.jsp"

driver = webdriver.Chrome(executable_path=r'/Users/admin/Desktop/chromedriver')

driver.get(url)

driver.find_element(By.NAME,'username').send_keys(username)

driver.find_element(By.NAME,'password').send_keys(password)

driver.find_elements(By.ID,"loginBtn")[0].click()

time.sleep(8) #Wait for login

driver.find_element(By.CLASS_NAME,"go-to-other-pages-buttons").click()

driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='schedule-item-container']/div[@class='schedule-item']").click()

Info = driver.find_element(By.CLASS_NAME, "schedule-table")

Data = Info.get_attribute('innerHTML')


All_shifts = []



time.sleep(3) #Let page load

emp = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][1]")

#Get employee hours from HS table & Save to All_shifts list

R1_Columns = 1

while R1_Columns != 54:
   for A in range (2,8):
        if (R1_Columns == 1):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row grayed border-top']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        elif (R1_Columns % 2 == 0):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        elif (R1_Columns != 53):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row grayed']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")                                 
        else:
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row grayed border-bottom']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        R1_shifts = (Row_data.text)
        R1_shifts = R1_shifts.replace("FOH General\n", "")
        All_shifts.append(R1_shifts)
   R1_Columns += 1


full_schedule = {}


days = dict({'Monday' : 'x', 'Tuesday' : 'x' , 'Wednesday' : 'x' , 'Thursday' : 'x' , 'Friday' : 'x' , 'Saturday' : 'x'})

#Get employee names from HS

for x in range (1,54):
   x = str (x)
   emp = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + x + "]/div[@class='schedule-table-cell employee']/div[@class='employee-name']")
   employee_id = emp.text
   full_schedule.update({employee_id : days})

#Create dataframe to display employee names and hours

df = pd.DataFrame(full_schedule).T 


#Replace letter x in dataframe with all the scheduled shifts

A = 0
B = 1
staff = 0

for p in range (0,53):
    nav = 0
    y = 1 
    for t in df.iloc[A:B]:
        df.iloc[A:B, nav:y] = All_shifts[staff]
        if nav == 6:
            nav == 0
            y == 1
            staff += 1
        else:
            nav += 1
            y += 1
            staff +=1
    A += 1
    B += 1

Today = ''
if 'Monday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,0:1].to_dict()
elif 'Tuesday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,1:2].to_dict() 
elif 'Wednesday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,2:3].to_dict()
elif 'Thursday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,3:4].to_dict()
elif 'Friday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,4:5].to_dict()
elif 'Saturday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:53,5:6].to_dict()



x,y,z = (DaySort(Today)) 
    
driver.close()





#Define Lists
employee_and_shift = []
employee_list = []
start_shift_list = []
end_shift_list = []
shift_length = []
time_2to3 = []
time_3to4 = []
time_4to5 = []
time_5to6 = []
time_6to7 = []
time_7to8 = []
time_8to9 = []
time_9to10 = []
FOHCrew = []

#Define positions as variables
#2-3
window2to3 = ()
DTbag2to3 = ()
Drinks2to3 = ()
Desserts2to3 = ()
Sauce2to3 = ()
FCbag2to3 = ()
Register1_2to3 = ()
Register2_2to3 = ()
Dining2to3 = ()
Cash2to3 = ()
Ipad1_2to3 = ()
Ipad2_2to3 = ()
Ipad3_2to3 = ()

#3-4
window3to4 = ()
DTbag3to4 = ()
Drinks3to4 = ()
Desserts3to4 = ()
Sauce3to4 = ()
FCbag3to4 = ()
Register1_3to4 = ()
Register2_3to4 = ()
Dining3to4 = ()
Cash3to4 = ()
Ipad1_3to4 = ()
Ipad2_3to4 = ()
Ipad3_3to4 = ()

#4-5
window4to5 = ()
DTbag4to5 = ()
Drinks4to5 = ()
Desserts4to5 = ()
Sauce4to5 = ()
FCbag4to5 = ()
Register1_4to5 = ()
Register2_4to5 = ()
Dining4to5 = ()
Cash4to5 = ()
Ipad1_4to5 = ()
Ipad2_4to5 = ()
Ipad3_4to5 = ()

#5-6
window5to6 = ()
DTbag5to6 = ()
Drinks5to6 = ()
Desserts5to6 = ()
Sauce5to6 = ()
FCbag5to6 = ()
Register1_5to6 = ()
Register2_5to6 = ()
Dining5to6 = ()
Cash5to6 = ()
Ipad1_5to6 = ()
Ipad2_5to6 = ()
Ipad3_5to6 = ()

#6-7
window6to7 = ()
DTbag6to7 = ()
Drinks6to7 = ()
Desserts6to7 = ()
Sauce6to7 = ()
FCbag6to7 = ()
Register1_6to7 = ()
Register2_6to7 = ()
Dining6to7 = ()
Cash6to7 = ()
Ipad1_6to7 = ()
Ipad2_6to7 = ()
Ipad3_6to7 = ()

#7-8
window7to8 = ()
DTbag7to8 = ()
Drinks7to8 = ()
Desserts7to8 = ()
Sauce7to8 = ()
FCbag7to8 = ()
Register1_7to8 = ()
Register2_7to8 = ()
Dining7to8 = ()
Cash7to8 = ()
Ipad1_7to8 = ()
Ipad2_7to8 = ()
Ipad3_7to8 = ()

#8-9
window8to9 = ()
DTbag8to9 = ()
Drinks8to9 = ()
Desserts8to9 = ()
Sauce8to9 = ()
FCbag8to9 = ()
Register1_8to9 = ()
Register2_8to9 = ()
Dining8to9 = ()
Cash8to9 = ()
Ipad1_8to9 = ()
Ipad2_8to9 = ()
Ipad3_8to9 = ()

#9-10
window9to10 = ()
DTbag9to10 = ()
Drinks9to10 = ()
Desserts9to10 = ()
Sauce9to10 = ()
FCbag9to10 = ()
Register1_9to10 = ()
Register2_9to10 = ()
Dining9to10 = ()

pos_Ipad1_4to5 = ()
pos_Ipad2_4to5 = ()
pos_Ipad3_4to5 = ()

employee_list = list(z)
start_shift = list(x)
end_shift = list(y) 

add = 0
for employee in employee_list:
    if (int (end_shift[add]) > 14 and int(start_shift[add]) <= 14):  
        time_2to3.append(employee)
    if (int (end_shift[add]) > 15 and int(start_shift[add]) <= 15):  
        time_3to4.append(employee) 
    if (int (end_shift[add]) > 16 and int(start_shift[add]) <= 16):  
        time_4to5.append(employee)
    if (int (end_shift[add]) > 17 and int(start_shift[add]) <= 17):  
        time_5to6.append(employee)
    if (int (end_shift[add]) > 18 and int(start_shift[add]) <= 18):  
        time_6to7.append(employee)
    if (int (end_shift[add]) > 19 and int(start_shift[add]) <= 19):  
        time_7to8.append(employee)
    if (int (end_shift[add]) > 20 and int(start_shift[add]) <= 20):  
        time_8to9.append(employee)
    if (int (end_shift[add]) > 21 and int(start_shift[add]) <= 21):  
        time_9to10.append(employee)       
    start_shift_list.append (start_shift)
    end_shift_list.append (end_shift)
    add += 1


#2-3
for x in time_2to3:
    #Register 1
        pos_Register1_2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Register1_2to3)
        if time_2to3 == []:
            continue          
    #Window    
        pos_window2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_window2to3)
        if time_2to3 == []:
            continue
    #DT bagging    
        pos_DTbag2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_DTbag2to3)
        if time_2to3 == []:
            continue
    #Drinks
        pos_Drinks2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Drinks2to3)
        if time_2to3 == []:
            continue
    #IPad 1
        pos_Ipad1_2to3 = ()
        pos_Ipad2_2to3 = ()
        pos_Ipad3_2to3 = ()
        pos_Ipad1_2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Ipad1_2to3)
        if time_2to3 == []:
            continue
    #Cash
        pos_Cash2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Cash2to3)
        if time_2to3 == []:
            continue
    #FC bagging
        pos_FCbag2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_FCbag2to3)
        if time_2to3 == []:
            continue
    #IPad 2
        pos_Ipad2_2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Ipad2_2to3)
        if time_2to3 == []:
            continue
    #Desserts
        pos_Desserts2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Desserts2to3)
        if time_2to3 == []:
            continue
    #Register 2
        pos_Register2_2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Register2_2to3)
        if time_2to3 == []:
            continue
    #IPad 3
        pos_Ipad3_2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Ipad3_2to3)
        if time_2to3 == []:
            continue
    #Dining Room
        pos_Dining2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Dining2to3)
        if time_2to3 == []:
            continue
    #Sauces
        pos_Sauce2to3 = random.choice(time_2to3)
        time_2to3.remove(pos_Sauce2to3)
        if time_2to3 == []:
            break
#3 to 4
for x in time_3to4:
    #Register 1
        pos_Register1_3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Register1_3to4)
        if time_3to4 == []:
            continue          
    #Window    
        pos_window3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_window3to4)
        if time_3to4 == []:
            continue
    #DT bagging    
        pos_DTbag3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_DTbag3to4)
        if time_3to4 == []:
            continue
    #Drinks
        pos_Drinks3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Drinks3to4)
        if time_3to4 == []:
            continue
    #IPad 1
        pos_Ipad1_3to4 = ()
        pos_Ipad2_3to4 = ()
        pos_Ipad3_3to4 = ()
        pos_Ipad1_3to4 = random.choice(time_3to4)
        while pos_Ipad1_3to4 == pos_Ipad1_2to3 or pos_Ipad1_3to4 == pos_Ipad2_2to3 or pos_Ipad1_3to4 == pos_Ipad3_2to3:
            if len(time_3to4) == 1:
                 pos_Ipad1_3to4 = ()   
                 break
            else:
                pos_Ipad1_3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Ipad1_3to4)
        if time_3to4 == []:
            continue
    #Cash
        pos_Cash3to4 = random.choice(time_3to4)
        if time_3to4 == []:
            continue
    #FC bagging
        pos_FCbag3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_FCbag3to4)
        if time_3to4 == []:
            continue
    #IPad 2
        pos_Ipad2_3to4 = random.choice(time_3to4)
        while pos_Ipad2_3to4 == pos_Ipad1_2to3 or pos_Ipad2_3to4 == pos_Ipad2_2to3 or pos_Ipad2_3to4 == pos_Ipad3_2to3:   
            if len(time_3to4) == 1:
                pos_Ipad2_3to4 = ()   
                break
            else:
                pos_Ipad2_3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Ipad2_3to4)
        if time_3to4 == []:
            continue
    #Desserts
        pos_Desserts3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Desserts3to4)
        if time_3to4 == []:
            continue
    #Register 2
        pos_Register2_3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Register2_3to4)
        if time_3to4 == []:
            continue
    #IPad 3
        pos_Ipad3_3to4 = random.choice(time_3to4)
        while pos_Ipad3_3to4 == pos_Ipad1_2to3 or pos_Ipad3_3to4 == pos_Ipad2_2to3 or pos_Ipad3_3to4 == pos_Ipad3_2to3:
                if len(time_3to4) == 1:
                    pos_Ipad3_3to4 = ()   
                    break
                else:
                    pos_Ipad3_3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Ipad3_3to4)
        if time_3to4 == []:
            continue
    #Dining Room
        pos_Dining3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Dining3to4)
        if time_3to4 == []:
            continue
    #Sauces
        pos_Sauce3to4 = random.choice(time_3to4)
        time_3to4.remove(pos_Sauce3to4)
        if time_3to4 == []:
            break
#4 to 5
for x in time_4to5:
    #Register 1
        pos_Register1_4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Register1_4to5)
        if time_4to5 == []:
            continue          
    #Window    
        pos_window4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_window4to5)
        if time_4to5 == []:
            continue
    #DT bagging    
        pos_DTbag4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_DTbag4to5)
        if time_4to5 == []:
            continue
    #Drinks
        pos_Drinks4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Drinks4to5)
        if time_4to5 == []:
            continue
    #IPad 1
        pos_Ipad1_4to5 = ()
        pos_Ipad2_4to5 = ()
        pos_Ipad3_4to5 = ()
        pos_Ipad1_4to5 = random.choice(time_4to5)
        while pos_Ipad1_4to5 == pos_Ipad1_3to4 or pos_Ipad1_4to5 == pos_Ipad2_3to4 or pos_Ipad1_4to5 == pos_Ipad3_3to4:
                if len(time_4to5) == 1:
                    pos_Ipad1_4to5 = ()   
                    break
                else:
                    pos_Ipad1_4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Ipad1_4to5)
        if time_4to5 == []:
            continue
    #Cash
        pos_Cash4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Cash4to5)
        if time_4to5 == []:
            continue
    #FC bagging
        pos_FCbag4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_FCbag4to5)
        if time_4to5 == []:
            continue
    #IPad 2
        pos_Ipad2_4to5 = random.choice(time_4to5)
        while pos_Ipad2_4to5 == pos_Ipad1_3to4 or pos_Ipad2_4to5 == pos_Ipad2_3to4 or pos_Ipad2_4to5 == pos_Ipad3_3to4:   
            if len(time_4to5) == 1:
                pos_Ipad2_4to5 = ()   
                break
            else:
                pos_Ipad2_4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Ipad2_4to5)
        if time_4to5 == []:
            continue
    #Desserts
        pos_Desserts4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Desserts4to5)
        if time_4to5 == []:
            continue
    #Register 2
        pos_Register2_4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Register2_4to5)
        if time_4to5 == []:
            continue
    #IPad 3
        pos_Ipad3_4to5 = random.choice(time_4to5)
        while pos_Ipad3_4to5 == pos_Ipad1_3to4 or pos_Ipad3_4to5 == pos_Ipad2_3to4 or pos_Ipad3_4to5 == pos_Ipad3_3to4:  
            if len(time_4to5) == 1:
                pos_Ipad3_4to5 = ()    
                break
            else:
                pos_Ipad3_4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Ipad3_4to5)
        if time_4to5 == []:
            continue
    #Dining Room
        pos_Dining4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Dining4to5)
        if time_4to5 == []:
            continue
    #Sauces
        pos_Sauce4to5 = random.choice(time_4to5)
        time_4to5.remove(pos_Sauce4to5)
        if time_4to5 == []:
            break
#5 to 6
for x in time_5to6:
    #Register 1
        pos_Register1_5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Register1_5to6)
        if time_5to6 == []:
            continue          
    #Window    
        pos_window5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_window5to6)
        if time_5to6 == []:
            continue
    #DT bagging    
        pos_DTbag5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_DTbag5to6)
        if time_5to6 == []:
            continue
    #Drinks
        pos_Drinks5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Drinks5to6)
        if time_5to6 == []:
            continue
    #IPad 1
        pos_Ipad1_5to6 = ()
        pos_Ipad2_5to6 = ()
        pos_Ipad3_5to6 = ()
        pos_Ipad1_5to6 = random.choice(time_5to6)
        while pos_Ipad1_5to6 == pos_Ipad1_4to5 or pos_Ipad1_5to6 == pos_Ipad2_4to5 or pos_Ipad1_5to6 == pos_Ipad3_4to5:
            if len(time_5to6) == 1:
                pos_Ipad1_5to6 = ()   
                break
            else:
                pos_Ipad1_5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Ipad1_5to6)
        if time_5to6 == []:
            continue
    #Cash
        pos_Cash5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Cash5to6)
        if time_5to6 == []:
            continue
    #FC bagging
        pos_FCbag5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_FCbag5to6)
        if time_5to6 == []:
            continue
    #IPad 2
        pos_Ipad2_5to6 = random.choice(time_5to6)
        while pos_Ipad2_5to6 == pos_Ipad1_4to5 or pos_Ipad2_5to6 == pos_Ipad2_4to5 or pos_Ipad2_5to6 == pos_Ipad3_4to5:   
            if len(time_5to6) == 1:
                pos_Ipad2_5to6 = ()   
                break
            else:
                pos_Ipad2_5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Ipad2_5to6)
        if time_5to6 == []:
            continue
    #Desserts
        pos_Desserts5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Desserts5to6)
        if time_5to6 == []:
            continue
    #Register 2
        pos_Register2_5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Register2_5to6)
        if time_5to6 == []:
            continue
    #IPad 3
        pos_Ipad3_5to6 = random.choice(time_5to6)
        while pos_Ipad3_5to6 == pos_Ipad1_4to5 or pos_Ipad3_5to6 == pos_Ipad2_4to5 or pos_Ipad3_5to6 == pos_Ipad3_4to5:
                if len(time_5to6) == 1:
                    pos_Ipad3_5to6 = ()   
                    break
                else:
                    pos_Ipad3_5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Ipad3_5to6)
        if time_5to6 == []:
            continue
    #Dining Room
        pos_Dining5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Dining5to6)
        if time_5to6 == []:
            continue
    #Sauces
        pos_Sauce5to6 = random.choice(time_5to6)
        time_5to6.remove(pos_Sauce5to6)
        if time_5to6 == []:
            break
#6 to 7
for x in time_6to7:
    #Register 1
        pos_Register1_6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Register1_6to7)
        if time_6to7 == []:
            continue          
    #Window    
        pos_window6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_window6to7)
        if time_6to7 == []:
            continue
    #DT bagging    
        pos_DTbag6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_DTbag6to7)
        if time_6to7 == []:
            continue
    #Drinks
        pos_Drinks6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Drinks6to7)
        if time_6to7 == []:
            continue
    #IPad 1
        pos_Ipad1_6to7 = ()
        pos_Ipad2_6to7 = ()
        pos_Ipad3_6to7 = ()
        pos_Ipad1_6to7 = random.choice(time_6to7)
        while pos_Ipad1_6to7 == pos_Ipad1_5to6 or pos_Ipad1_6to7 == pos_Ipad2_5to6 or pos_Ipad1_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad1_6to7 = ()   
                break
            else:
                pos_Ipad1_6to7 = random.choice(time_6to7)           
        time_6to7.remove(pos_Ipad1_6to7)
        if time_6to7 == []:
            continue
    #Cash
        pos_Cash6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Cash6to7)
        if time_6to7 == []:
            continue
    #FC bagging
        pos_FCbag6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_FCbag6to7)
        if time_6to7 == []:
            continue
    #IPad 2
        pos_Ipad2_6to7 = random.choice(time_6to7)
        while pos_Ipad2_6to7 == pos_Ipad1_5to6 or pos_Ipad2_6to7 == pos_Ipad2_5to6 or pos_Ipad2_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad2_6to7 = ()   
                break
            else:
                pos_Ipad2_6to7 = random.choice(time_6to7)   
        time_6to7.remove(pos_Ipad2_6to7)
        if time_6to7 == []:
            continue
    #Desserts
        pos_Desserts6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Desserts6to7)
        if time_6to7 == []:
            continue
    #Register 2
        pos_Register2_6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Register2_6to7)
        if time_6to7 == []:
            continue
    #IPad 3
        pos_Ipad3_6to7 = random.choice(time_6to7)
        while pos_Ipad3_6to7 == pos_Ipad1_5to6 or pos_Ipad3_6to7 == pos_Ipad2_5to6 or pos_Ipad3_6to7 == pos_Ipad3_5to6:
            if len(time_6to7) == 1:
                pos_Ipad3_6to7 = ()   
                break   
            else:
                pos_Ipad3_6to7 = random.choice(time_6to7)     
        time_6to7.remove(pos_Ipad3_6to7)
        if time_6to7 == []:
            continue
    #Dining Room
        pos_Dining6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Dining6to7)
        if time_6to7 == []:
            continue
    #Sauces
        pos_Sauce6to7 = random.choice(time_6to7)
        time_6to7.remove(pos_Sauce6to7)
        if time_6to7 == []:
            break
        
#7 to 8
for x in time_7to8:
    #Register 1
        pos_Register1_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Register1_7to8)
        if time_7to8 == []:
            continue          
    #Window    
        pos_window7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_window7to8)
        if time_7to8 == []:
            continue
    #DT bagging    
        pos_DTbag7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_DTbag7to8)
        if time_7to8 == []:
            continue
    #Drinks
        pos_Drinks7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Drinks7to8)
        if time_7to8 == []:
            continue
    #IPad 1
        pos_Ipad1_7to8 = ()
        pos_Ipad2_7to8 = ()
        pos_Ipad3_7to8 = ()
        pos_Ipad1_7to8 = random.choice(time_7to8)
        while pos_Ipad1_7to8 == pos_Ipad1_6to7 or pos_Ipad1_7to8 == pos_Ipad2_6to7 or pos_Ipad1_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad1_7to8 = ()   
                break   
            else:
                pos_Ipad1_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Ipad1_7to8)
        if time_7to8 == []:
            continue
    #Cash
        pos_Cash7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Cash7to8)
        if time_7to8 == []:
            continue
    #Register 1
        pos_Register1_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Register1_7to8)
        if time_7to8 == []:
             continue
    #FC bagging
        pos_FCbag7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_FCbag7to8)
        if time_7to8 == []:
            continue
    #IPad 2
        pos_Ipad2_7to8 = random.choice(time_7to8)
        while pos_Ipad2_7to8 == pos_Ipad1_6to7 or pos_Ipad2_7to8 == pos_Ipad2_6to7 or pos_Ipad2_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad2_7to8 = ()   
                break   
            else:
                pos_Ipad2_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Ipad2_7to8)
        if time_7to8 == []:
            continue
    #Desserts
        pos_Desserts7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Desserts7to8)
        if time_7to8 == []:
            continue
    #Register 2
        pos_Register2_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Register2_7to8)
        if time_7to8 == []:
            continue
    #IPad 3
        pos_Ipad3_7to8 = random.choice(time_7to8)
        while pos_Ipad3_7to8 == pos_Ipad1_6to7 or pos_Ipad3_7to8 == pos_Ipad2_6to7 or pos_Ipad3_7to8 == pos_Ipad3_6to7:
            if len(time_7to8) == 1:
                pos_Ipad3_7to8 = ()   
                break   
            else:
                pos_Ipad3_7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Ipad3_7to8)
        if time_7to8 == []:
            continue
    #Dining Room
        pos_Dining7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Dining7to8)
        if time_7to8 == []:
            continue
    #Sauces
        pos_Sauce7to8 = random.choice(time_7to8)
        time_7to8.remove(pos_Sauce7to8)
        if time_7to8 == []:
            break
#8 to 9            
for x in time_8to9:
#Window    
        pos_window8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_window8to9)
        if time_8to9 == []:
            continue
    #DT bagging    
        pos_DTbag8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_DTbag8to9)
        if time_8to9 == []:
            continue
    #Drinks
        pos_Drinks8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Drinks8to9)
        if time_8to9 == []:
            continue
    #Dining Room
        pos_Dining8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Dining8to9)
        if time_8to9 == []:
            continue        
    #IPad 1
        pos_Ipad1_8to9 = ()
        pos_Ipad2_8to9 = ()
        pos_Ipad3_8to9 = ()
        pos_Ipad1_8to9 = random.choice(time_8to9)
        while pos_Ipad1_8to9 == pos_Ipad1_7to8 or pos_Ipad1_8to9 == pos_Ipad2_7to8 or pos_Ipad1_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad1_8to9 = ()   
                break  
            else:
                pos_Ipad1_8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Ipad1_8to9)
        if time_8to9 == []:
            continue
    #Cash
        pos_Cash8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Cash8to9)
        if time_8to9 == []:
            continue
    #Register 1
        pos_Register1_8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Register1_8to9)
        if time_8to9 == []:
             continue
    #FC bagging
        pos_FCbag8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_FCbag8to9)
        if time_8to9 == []:
            continue
    #IPad 2
        pos_Ipad2_8to9 = random.choice(time_8to9)
        while pos_Ipad2_8to9 == pos_Ipad1_7to8 or pos_Ipad2_8to9 == pos_Ipad2_7to8 or pos_Ipad2_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad2_8to9 = ()   
                break  
            else:
                pos_Ipad2_8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Ipad2_8to9)
        if time_8to9 == []:
            continue
    #Desserts
        pos_Desserts8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Desserts8to9)
        if time_8to9 == []:
            continue
    #Register 2
        pos_Register2_8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Register2_8to9)
        if time_8to9 == []:
            continue
    #IPad 3
        pos_Ipad3_8to9 = random.choice(time_8to9)
        while pos_Ipad3_8to9 == pos_Ipad1_7to8 or pos_Ipad3_8to9 == pos_Ipad2_7to8 or pos_Ipad3_8to9 == pos_Ipad3_7to8:
            if len(time_8to9) == 1:
                pos_Ipad3_8to9 = ()   
                break  
            else:
                pos_Ipad3_8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Ipad3_8to9)
        if time_8to9 == []:
            continue
    #Sauces
        pos_Sauce8to9 = random.choice(time_8to9)
        time_8to9.remove(pos_Sauce8to9)
        if time_8to9 == []:
            break        
#9 to 10            
for x in time_9to10:
     #Window    
        pos_window9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_window9to10)
        if time_9to10 == []:
            continue
    #DT bagging    
        pos_DTbag9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_DTbag9to10)
        if time_9to10 == []:
            continue
    #Drinks
        pos_Drinks9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_Drinks9to10)
        if time_9to10 == []:
            continue
    #Dining Room
        pos_Dining9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_Dining9to10)
        if time_9to10 == []:
            continue
    #Register 1
        pos_Register1_9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_Register1_9to10)
        if time_9to10 == []:
             continue
    #FC bagging
        pos_FCbag9to10 = random.choice(time_9to10)
        time_9to10.remove(pos_FCbag9to10)
        if time_9to10 == []:
            continue

#Output as word document

def iter_paragraphs_of_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                  yield paragraph
                yield from iter_paragraphs_of_tables(cell.tables)

for paragraph in iter_paragraphs_of_tables(doc.tables):
     if '{{Win2to3}}' in paragraph.text:
      if window2to3 != ():
       paragraph.text = paragraph.text.replace("{{Win2to3}}", window2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Win2to3}}", "")   
     if '{{DTBag2to3}}' in paragraph.text:
      if DTbag2to3 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", DTbag2to3 + " 2-3")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", "")
     if '{{Drinks2to3}}' in paragraph.text:
      if Drinks2to3 != ():
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", Drinks2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", "")   
     if '{{Desserts2to3}}' in paragraph.text:
        if Desserts2to3 != ():
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", Desserts2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", "")   
     if '{{Sauce2to3}}' in paragraph.text:
        if Sauce2to3 != ():
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", Sauce2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", "")   
     if '{{FCBag2to3}}' in paragraph.text:
        if FCbag2to3 != ():
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", FCbag2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", "")   
     if '{{Reg1_2to3}}' in paragraph.text:
        if Register1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", Register1_2to3 + " 2-3")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", "") 
     if '{{Reg2_2to3}}' in paragraph.text:
        if Register2_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", Register2_2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", "")  
     if '{{Dining_2to3}}' in paragraph.text:
        if Dining2to3 != ():
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", Dining2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", "")
     if '{{Cash2to3}}' in paragraph.text:
        if Cash2to3 != ():
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", Cash2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", "")   
     if '{{IPad1_2to3}}' in paragraph.text:
        if Ipad1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", Ipad1_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", "")    
     if '{{IPad2_2to3}}' in paragraph.text:
        if Ipad2_2to3 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", Ipad2_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", "")   
     if '{{IPad3_2to3}}' in paragraph.text:
        if Ipad3_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", Ipad3_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", "")
     if '{{Win3to4}}' in paragraph.text:
      if window3to4 != ():
       paragraph.text = paragraph.text.replace("{{Win3to4}}", window3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Win3to4}}", "")   
     if '{{DTBag3to4}}' in paragraph.text:
      if DTbag3to4 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", DTbag3to4 + " 3-4")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", "")
     if '{{Drinks3to4}}' in paragraph.text:
      if Drinks3to4 != ():
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", Drinks3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", "")   
     if '{{Desserts3to4}}' in paragraph.text:
        if Desserts3to4 != ():
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", Desserts3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", "")   
     if '{{Sauce3to4}}' in paragraph.text:
        if Sauce3to4 != ():
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", Sauce3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", "")   
     if '{{FCBag3to4}}' in paragraph.text:
        if FCbag3to4 != ():
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", FCbag3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", "")   
     if '{{Reg1_3to4}}' in paragraph.text:
        if Register1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", Register1_3to4 + " 3-4")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", "") 
     if '{{Reg2_3to4}}' in paragraph.text:
        if Register2_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", Register2_3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", "")  
     if '{{Dining_3to4}}' in paragraph.text:
        if Dining3to4 != ():
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", Dining3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", "")
     if '{{Cash3to4}}' in paragraph.text:
        if Cash3to4 != ():
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", Cash3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", "")   
     if '{{IPad1_3to4}}' in paragraph.text:
        if Ipad1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", Ipad1_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", "")    
     if '{{IPad2_3to4}}' in paragraph.text:
        if Ipad2_3to4 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", Ipad2_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", "")   
     if '{{IPad3_3to4}}' in paragraph.text:
        if Ipad3_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", Ipad3_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", "")       
     if '{{Win4to5}}' in paragraph.text:
      if window4to5 != ():
       paragraph.text = paragraph.text.replace("{{Win4to5}}", window4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Win4to5}}", "")   
     if '{{DTBag4to5}}' in paragraph.text:
      if DTbag4to5 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", DTbag4to5 + " 4-5")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", "")
     if '{{Drinks4to5}}' in paragraph.text:
      if Drinks4to5 != ():
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", Drinks4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", "")   
     if '{{Desserts4to5}}' in paragraph.text:
        if Desserts4to5 != ():
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", Desserts4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", "")   
     if '{{Sauce4to5}}' in paragraph.text:
        if Sauce4to5 != ():
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", Sauce4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", "")   
     if '{{FCBag4to5}}' in paragraph.text:
        if FCbag4to5 != ():
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", FCbag4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", "")   
     if '{{Reg1_4to5}}' in paragraph.text:
        if Register1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", Register1_4to5 + " 4-5")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", "") 
     if '{{Reg2_4to5}}' in paragraph.text:
        if Register2_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", Register2_4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", "")  
     if '{{Dining_4to5}}' in paragraph.text:
        if Dining4to5 != ():
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", Dining4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", "")
     if '{{Cash4to5}}' in paragraph.text:
        if Cash4to5 != ():
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", Cash4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", "")   
     if '{{IPad1_4to5}}' in paragraph.text:
        if Ipad1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", Ipad1_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", "")    
     if '{{IPad2_4to5}}' in paragraph.text:
        if Ipad2_4to5 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", Ipad2_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", "")   
     if '{{IPad3_4to5}}' in paragraph.text:
        if Ipad3_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", Ipad3_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", "")
     if '{{Win5to6}}' in paragraph.text:
      if window5to6 != ():
       paragraph.text = paragraph.text.replace("{{Win5to6}}", window5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Win5to6}}", "")   
     if '{{DTBag5to6}}' in paragraph.text:
      if DTbag5to6 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", DTbag5to6 + " 5-6")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", "")
     if '{{Drinks5to6}}' in paragraph.text:
      if Drinks5to6 != ():
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", Drinks5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", "")   
     if '{{Desserts5to6}}' in paragraph.text:
        if Desserts5to6 != ():
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", Desserts5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", "")   
     if '{{Sauce5to6}}' in paragraph.text:
        if Sauce5to6 != ():
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", Sauce5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", "")   
     if '{{FCBag5to6}}' in paragraph.text:
        if FCbag5to6 != ():
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", FCbag5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", "")   
     if '{{Reg1_5to6}}' in paragraph.text:
        if Register1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", Register1_5to6 + " 5-6")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", "") 
     if '{{Reg2_5to6}}' in paragraph.text:
        if Register2_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", Register2_5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", "")  
     if '{{Dining_5to6}}' in paragraph.text:
        if Dining5to6 != ():
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", Dining5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", "")
     if '{{Cash5to6}}' in paragraph.text:
        if Cash5to6 != ():
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", Cash5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", "")   
     if '{{IPad1_5to6}}' in paragraph.text:
        if Ipad1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", Ipad1_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", "")    
     if '{{IPad2_5to6}}' in paragraph.text:
        if Ipad2_5to6 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", Ipad2_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", "")   
     if '{{IPad3_5to6}}' in paragraph.text:
        if Ipad3_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", Ipad3_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", "")         
     if '{{Win6to7}}' in paragraph.text:
      if window6to7 != ():
       paragraph.text = paragraph.text.replace("{{Win6to7}}", window6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Win6to7}}", "")   
     if '{{DTBag6to7}}' in paragraph.text:
      if DTbag6to7 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", DTbag6to7 + " 6-7")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", "")
     if '{{Drinks6to7}}' in paragraph.text:
      if Drinks6to7 != ():
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", Drinks6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", "")   
     if '{{Desserts6to7}}' in paragraph.text:
        if Desserts6to7 != ():
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", Desserts6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", "")   
     if '{{Sauce6to7}}' in paragraph.text:
        if Sauce6to7 != ():
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", Sauce6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", "")   
     if '{{FCBag6to7}}' in paragraph.text:
        if FCbag6to7 != ():
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", FCbag6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", "")   
     if '{{Reg1_6to7}}' in paragraph.text:
        if Register1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", Register1_6to7 + " 6-7")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", "") 
     if '{{Reg2_6to7}}' in paragraph.text:
        if Register2_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", Register2_6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", "")  
     if '{{Dining_6to7}}' in paragraph.text:
        if Dining6to7 != ():
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", Dining6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", "")
     if '{{Cash6to7}}' in paragraph.text:
        if Cash6to7 != ():
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", Cash6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", "")   
     if '{{IPad1_6to7}}' in paragraph.text:
        if Ipad1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", Ipad1_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", "")    
     if '{{IPad2_6to7}}' in paragraph.text:
        if Ipad2_6to7 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", Ipad2_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", "")   
     if '{{IPad3_6to7}}' in paragraph.text:
        if Ipad3_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", Ipad3_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", "")    
     if '{{Win7to8}}' in paragraph.text:
      if window7to8 != ():
       paragraph.text = paragraph.text.replace("{{Win7to8}}", window7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Win7to8}}", "")   
     if '{{DTBag7to8}}' in paragraph.text:
      if DTbag7to8 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", DTbag7to8 + " 7-8")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", "")
     if '{{Drinks7to8}}' in paragraph.text:
      if Drinks7to8 != ():
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", Drinks7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", "")   
     if '{{Desserts7to8}}' in paragraph.text:
        if Desserts7to8 != ():
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", Desserts7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", "")   
     if '{{Sauce7to8}}' in paragraph.text:
        if Sauce7to8 != ():
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", Sauce7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", "")   
     if '{{FCBag7to8}}' in paragraph.text:
        if FCbag7to8 != ():
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", FCbag7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", "")   
     if '{{Reg1_7to8}}' in paragraph.text:
        if Register1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", Register1_7to8 + " 7-8")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", "") 
     if '{{Reg2_7to8}}' in paragraph.text:
        if Register2_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", Register2_7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", "")  
     if '{{Dining_7to8}}' in paragraph.text:
        if Dining7to8 != ():
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", Dining7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", "")
     if '{{Cash7to8}}' in paragraph.text:
        if Cash7to8 != ():
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", Cash7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", "")   
     if '{{IPad1_7to8}}' in paragraph.text:
        if Ipad1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", Ipad1_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", "")    
     if '{{IPad2_7to8}}' in paragraph.text:
        if Ipad2_7to8 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", Ipad2_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", "")   
     if '{{IPad3_7to8}}' in paragraph.text:
        if Ipad3_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", Ipad3_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", "")
     if '{{Win8to9}}' in paragraph.text:
      if window8to9 != ():
       paragraph.text = paragraph.text.replace("{{Win8to9}}", window8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Win8to9}}", "")   
     if '{{DTBag8to9}}' in paragraph.text:
      if DTbag8to9 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", DTbag8to9 + " 8-9")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", "")
     if '{{Drinks8to9}}' in paragraph.text:
      if Drinks8to9 != ():
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", Drinks8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", "")   
     if '{{Desserts8to9}}' in paragraph.text:
        if Desserts8to9 != ():
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", Desserts8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", "")   
     if '{{Sauce8to9}}' in paragraph.text:
        if Sauce8to9 != ():
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", Sauce8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", "")   
     if '{{FCBag8to9}}' in paragraph.text:
        if FCbag8to9 != ():
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", FCbag8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", "")   
     if '{{Reg1_8to9}}' in paragraph.text:
        if Register1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", Register1_8to9 + " 8-9")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", "") 
     if '{{Reg2_8to9}}' in paragraph.text:
        if Register2_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", Register2_8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", "")  
     if '{{Dining_8to9}}' in paragraph.text:
        if Dining8to9 != ():
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", Dining8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", "")
     if '{{Cash8to9}}' in paragraph.text:
        if Cash8to9 != ():
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", Cash8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", "")   
     if '{{IPad1_8to9}}' in paragraph.text:
        if Ipad1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", Ipad1_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", "")    
     if '{{IPad2_8to9}}' in paragraph.text:
        if Ipad2_8to9 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", Ipad2_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", "")   
     if '{{IPad3_8to9}}' in paragraph.text:
        if Ipad3_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", Ipad3_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", "")         
     if '{{Win9to10}}' in paragraph.text:
      if window9to10 != ():
       paragraph.text = paragraph.text.replace("{{Win9to10}}", window9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Win9to10}}", "")   
     if '{{DTBag9to10}}' in paragraph.text:
      if DTbag9to10 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", DTbag9to10 + " 9-10")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", "")
     if '{{Drinks9to10}}' in paragraph.text:
      if Drinks9to10 != ():
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", Drinks9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", "")   
     if '{{Desserts9to10}}' in paragraph.text:
        if Desserts9to10 != ():
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", Desserts9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", "")   
     if '{{Sauce9to10}}' in paragraph.text:
        if Sauce9to10 != ():
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", Sauce9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", "")   
     if '{{FCBag9to10}}' in paragraph.text:
        if FCbag9to10 != ():
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", FCbag9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", "")   
     if '{{Reg1_9to10}}' in paragraph.text:
        if Register1_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", Register1_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", "") 
     if '{{Reg2_9to10}}' in paragraph.text:
        if Register2_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", Register2_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", "")  
     if '{{Dining_9to10}}' in paragraph.text:
        if Dining9to10 != ():
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", Dining9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", "")
     for y in range (0,len (FOHCrew)):
         team_order = '{{FOHCrew'+ str (y) + '}}'
         if team_order in paragraph.text:
             paragraph.text = paragraph.text.replace(team_order, FOHCrew[y])
if '{{Date}}' in paragraph.text:
  paragraph.text = paragraph.text.replace("{{Date}}", print_date)            
font = paragraph.style.font         
font.size = Pt(9)         
doc.save('ShiftPlanner.docx')         
os.startfile('/Users/admin/Desktop/ShiftPlanner.docx')

   
                              