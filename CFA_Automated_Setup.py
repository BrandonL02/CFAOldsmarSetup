import os

import datetime as dt


import calendar

import random

import pandas as pd

import time

from selenium import webdriver

from selenium.webdriver.common.by import By

#Open document for editing
from docx import Document
from docx.shared import Pt 

doc = Document('C:/Users/Brand/Downloads/Front Counter Setup Code.docx')

tables = doc.tables 
#Get date and weekday within python
date = dt.date.today() #+ dt.timedelta(days=-1) <- To change to yesterday

print_date = (calendar.day_name[date.weekday()] + ', ' + date.strftime("%B %d"))

Workers = []
Begin = []
End = []
Duration = []


def minute_offset(user_var):
    if user_var >= 30:
        user_var = 1
    else:
        user_var = 0    
    return int(user_var)    

def time_to_list(Enter):
    Hours_int = []
    for e in Enter:
        e = time.strptime(e,'%I:%M%p')
        if 0 in e[:5]:
           # if 'PM' in e:
            #   e = e[3] + 12
            e = e[3]    
        else:
            #if 'PM' in e:
             #   estimate = minute_offset(e[4])
            #    e = e[3] + estimate + 12
            estimate = minute_offset(e[4])
            e = e[3] + estimate    
        Hours_int.append(e)
    return Hours_int
    
    

def Separate(info):
    for key, value in dict(info).items():
        Workers.append(key)
        Duration.append(value)
    for x in Duration:
        on, off = x.split(' - ', 2)
        Begin.append(on)
        End.append(off)
    x = time_to_list(Begin)
    y = time_to_list(End)
    z = Workers
    return x,y,z
    
def remove_unscheduled(Day):
    for key, value in dict(Day).items():
        if value == '':
            del Day[key] 
            
def DaySort(Weekday):
    for h in Weekday:
        Working = (Weekday[h])
        remove_unscheduled(Working)
        x,y,z = Separate(Working)
    return x,y,z
     

#Access Hotschedules with credentials and navigate pages

username = '2445922'

password = 'Monkey2002'

url = "https://www.hotschedules.com/hs/login.jsp"

driver = webdriver.Chrome(executable_path=r'C:/Users/Brand\Desktop/chromedriver.exe')

driver.get(url)

driver.find_element(By.NAME,'username').send_keys(username)

driver.find_element(By.NAME,'password').send_keys(password)

driver.find_elements(By.ID,"loginBtn")[0].click()

time.sleep(2) #Wait for login

driver.find_element(By.CLASS_NAME,"go-to-other-pages-buttons").click()

driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='schedule-item-container']/div[@class='schedule-item']").click()

Info = driver.find_element(By.CLASS_NAME, "schedule-table")

Data = Info.get_attribute('innerHTML')


All_shifts = []



time.sleep(3) #Let page load

emp = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][1]")

#Get employee hours from HS table & Save to All_shifts list

R1_Columns = 1

while R1_Columns != 53:
   for A in range (2,8):
        if (R1_Columns == 1):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row grayed border-top']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        elif (R1_Columns % 2 == 0 and R1_Columns != 52):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        elif (R1_Columns != 52):
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row grayed']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")                                 
        else:
            Row_data = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + str (R1_Columns) + "]/div[@class='days-row border-bottom']/div[@class='schedule-table-cell scheduled-shift']["  + str (A) + "]")
        R1_shifts = (Row_data.text)
        R1_shifts = R1_shifts.replace("FOH General\n", "")
        All_shifts.append(R1_shifts)
   R1_Columns += 1
   if R1_Columns == 53:
       break


full_schedule = {}


days = dict({'Monday' : 'x', 'Tuesday' : 'x' , 'Wednesday' : 'x' , 'Thursday' : 'x' , 'Friday' : 'x' , 'Saturday' : 'x'})

#Get employee names from HS

for x in range (1,53):
   x = str (x)
   emp = driver.find_element(By.XPATH, "/html[@class='hs-echo-ui clarifi-lobal-nav']/body[@class='screen-employeeHome employee-home-page']/div[@id='entire-taskarea-wrapper']/div[@id='root']/div[@class='employee-home']/div/div[@class='employee-home-inner-page']/div[@class='page-content']/div[@class='emp-home-widgets-dashboard']/div[@class='inner-content widgets-count-3']/div[@class='emp-home-base-widget is-expanded']/div[@class='inner different-layout']/div[@class='expanded-view all-schedules']/div[@class='children-container']/div[@class='emp-home-all-schedules-widget']/div[@class='all-schedules-container']/div[@class='all-schedules-details']/div[@class='my-schedules-items']/div[2]/div[@class='full-schedule']/div[@class='schedule-table']/div[@class='schedule-table-row'][" + x + "]/div[@class='schedule-table-cell employee']/div[@class='employee-name']")
   employee_id = emp.text
   full_schedule.update({employee_id : days})

#Create dataframe to display employee names and hours

df = pd.DataFrame(full_schedule).T 


#Replace letter x in dataframe with all the scheduled shifts

A = 0
B = 1
staff = 0

for p in range (0,52):
    nav = 0
    y = 1 
    for t in df.iloc[A:B]:
        df.iloc[A:B, nav:y] = All_shifts[staff]
        if nav == 6:
            nav == 0
            y == 1
            staff += 1
        else:
            nav += 1
            y += 1
            staff +=1
    A += 1
    B += 1

Today = ''

if 'Monday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,0:1].to_dict()
elif 'Tuesday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,1:2].to_dict() 
elif 'Wednesday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,2:3].to_dict()
elif 'Thursday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,3:4].to_dict()
elif 'Friday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,4:5].to_dict()
elif 'Saturday' in (calendar.day_name[date.weekday()]):
    Today = df.iloc[0:52,5:6].to_dict()



x,y,z = (DaySort(Today)) 
    
driver.close()





#Define Lists
employee_and_shift = []
employee_list = []
start_shift_list = []
end_shift_list = []
shift_length = []
time_2to3 = []
time_3to4 = []
time_4to5 = []
time_5to6 = []
time_6to7 = []
time_7to8 = []
time_8to9 = []
time_9to10 = []
FOHCrew = []

#Define positions as variables
#2-3
window2to3 = ()
DTbag2to3 = ()
Drinks2to3 = ()
Desserts2to3 = ()
Sauce2to3 = ()
FCbag2to3 = ()
Register1_2to3 = ()
Register2_2to3 = ()
Dining2to3 = ()
Cash2to3 = ()
Ipad1_2to3 = ()
Ipad2_2to3 = ()
Ipad3_2to3 = ()

#3-4
window3to4 = ()
DTbag3to4 = ()
Drinks3to4 = ()
Desserts3to4 = ()
Sauce3to4 = ()
FCbag3to4 = ()
Register1_3to4 = ()
Register2_3to4 = ()
Dining3to4 = ()
Cash3to4 = ()
Ipad1_3to4 = ()
Ipad2_3to4 = ()
Ipad3_3to4 = ()

#4-5
window4to5 = ()
DTbag4to5 = ()
Drinks4to5 = ()
Desserts4to5 = ()
Sauce4to5 = ()
FCbag4to5 = ()
Register1_4to5 = ()
Register2_4to5 = ()
Dining4to5 = ()
Cash4to5 = ()
Ipad1_4to5 = ()
Ipad2_4to5 = ()
Ipad3_4to5 = ()

#5-6
window5to6 = ()
DTbag5to6 = ()
Drinks5to6 = ()
Desserts5to6 = ()
Sauce5to6 = ()
FCbag5to6 = ()
Register1_5to6 = ()
Register2_5to6 = ()
Dining5to6 = ()
Cash5to6 = ()
Ipad1_5to6 = ()
Ipad2_5to6 = ()
Ipad3_5to6 = ()

#6-7
window6to7 = ()
DTbag6to7 = ()
Drinks6to7 = ()
Desserts6to7 = ()
Sauce6to7 = ()
FCbag6to7 = ()
Register1_6to7 = ()
Register2_6to7 = ()
Dining6to7 = ()
Cash6to7 = ()
Ipad1_6to7 = ()
Ipad2_6to7 = ()
Ipad3_6to7 = ()

#7-8
window7to8 = ()
DTbag7to8 = ()
Drinks7to8 = ()
Desserts7to8 = ()
Sauce7to8 = ()
FCbag7to8 = ()
Register1_7to8 = ()
Register2_7to8 = ()
Dining7to8 = ()
Cash7to8 = ()
Ipad1_7to8 = ()
Ipad2_7to8 = ()
Ipad3_7to8 = ()

#8-9
window8to9 = ()
DTbag8to9 = ()
Drinks8to9 = ()
Desserts8to9 = ()
Sauce8to9 = ()
FCbag8to9 = ()
Register1_8to9 = ()
Register2_8to9 = ()
Dining8to9 = ()
Cash8to9 = ()
Ipad1_8to9 = ()
Ipad2_8to9 = ()
Ipad3_8to9 = ()

#9-10
window9to10 = ()
DTbag9to10 = ()
Drinks9to10 = ()
Desserts9to10 = ()
Sauce9to10 = ()
FCbag9to10 = ()
Register1_9to10 = ()
Register2_9to10 = ()
Dining9to10 = ()

Ipad1_4to5 = ()
Ipad2_4to5 = ()
Ipad3_4to5 = ()

employee_list = list(z)
Name_cut = []
start_shift = list(x)
end_shift = list(y) 

for x in employee_list:
    sep = ' '
    x = head,sep,tail = x.partition(sep)
    tail = (tail[0:1] + '.')
    x = ''.join(head + ' ' + tail)
    Name_cut.append(x)

add = 0
for employee in Name_cut:
    if (int (end_shift[add]) > 14 and int(start_shift[add]) <= 14):  
        time_2to3.append(employee)
    if (int (end_shift[add]) > 15 and int(start_shift[add]) <= 15):  
        time_3to4.append(employee) 
    if (int (end_shift[add]) > 16 and int(start_shift[add]) <= 16):  
        time_4to5.append(employee)
    if (int (end_shift[add]) > 17 and int(start_shift[add]) <= 17):  
        time_5to6.append(employee)
    if (int (end_shift[add]) > 18 and int(start_shift[add]) <= 18):  
        time_6to7.append(employee)
    if (int (end_shift[add]) > 19 and int(start_shift[add]) <= 19):  
        time_7to8.append(employee)
    if (int (end_shift[add]) > 20 and int(start_shift[add]) <= 20):  
        time_8to9.append(employee)
    if (int (end_shift[add]) > 21 and int(start_shift[add]) <= 21):  
        time_9to10.append(employee)
    FOHCrew.append(employee)    
    add += 1


#2-3
for x in time_2to3:
    #Register 1
        Register1_2to3 = random.choice(time_2to3)
        time_2to3.remove(Register1_2to3)
        if time_2to3 == []:
            continue          
    #Window    
        window2to3 = random.choice(time_2to3)
        time_2to3.remove(window2to3)
        if time_2to3 == []:
            continue
    #DT bagging    
        DTbag2to3 = random.choice(time_2to3)
        time_2to3.remove(DTbag2to3)
        if time_2to3 == []:
            continue
    #Drinks
        Drinks2to3 = random.choice(time_2to3)
        time_2to3.remove(Drinks2to3)
        if time_2to3 == []:
            continue
    #IPad 1
        Ipad1_2to3 = ()
        Ipad2_2to3 = ()
        Ipad3_2to3 = ()
        Ipad1_2to3 = random.choice(time_2to3)
        time_2to3.remove(Ipad1_2to3)
        if time_2to3 == []:
            continue
    #Cash
        Cash2to3 = random.choice(time_2to3)
        time_2to3.remove(Cash2to3)
        if time_2to3 == []:
            continue
    #FC bagging
        FCbag2to3 = random.choice(time_2to3)
        time_2to3.remove(FCbag2to3)
        if time_2to3 == []:
            continue
    #IPad 2
        Ipad2_2to3 = random.choice(time_2to3)
        time_2to3.remove(Ipad2_2to3)
        if time_2to3 == []:
            continue
    #Desserts
        Desserts2to3 = random.choice(time_2to3)
        time_2to3.remove(Desserts2to3)
        if time_2to3 == []:
            continue
    #Register 2
        Register2_2to3 = random.choice(time_2to3)
        time_2to3.remove(Register2_2to3)
        if time_2to3 == []:
            continue
    #IPad 3
        Ipad3_2to3 = random.choice(time_2to3)
        time_2to3.remove(Ipad3_2to3)
        if time_2to3 == []:
            continue
    #Dining Room
        Dining2to3 = random.choice(time_2to3)
        time_2to3.remove(Dining2to3)
        if time_2to3 == []:
            continue
    #Sauces
        Sauce2to3 = random.choice(time_2to3)
        time_2to3.remove(Sauce2to3)
        if time_2to3 == []:
            break
#3 to 4
for x in time_3to4:
    #Register 1
        Register1_3to4 = random.choice(time_3to4)
        time_3to4.remove(Register1_3to4)
        if time_3to4 == []:
            continue          
    #Window    
        window3to4 = random.choice(time_3to4)
        time_3to4.remove(window3to4)
        if time_3to4 == []:
            continue
    #DT bagging    
        DTbag3to4 = random.choice(time_3to4)
        time_3to4.remove(DTbag3to4)
        if time_3to4 == []:
            continue
    #Drinks
        Drinks3to4 = random.choice(time_3to4)
        time_3to4.remove(Drinks3to4)
        if time_3to4 == []:
            continue
    #IPad 1
        Ipad1_3to4 = ()
        Ipad2_3to4 = ()
        Ipad3_3to4 = ()
        Ipad1_3to4 = random.choice(time_3to4)
        while Ipad1_3to4 == Ipad1_2to3 or Ipad1_3to4 == Ipad2_2to3 or Ipad1_3to4 == Ipad3_2to3:
            if len(time_3to4) == 1:
                 Ipad1_3to4 = ()   
                 break
            else:
                Ipad1_3to4 = random.choice(time_3to4)
        time_3to4.remove(Ipad1_3to4)
        if time_3to4 == []:
            continue
    #Cash
        Cash3to4 = random.choice(time_3to4)
        if time_3to4 == []:
            continue
    #FC bagging
        FCbag3to4 = random.choice(time_3to4)
        time_3to4.remove(FCbag3to4)
        if time_3to4 == []:
            continue
    #IPad 2
        Ipad2_3to4 = random.choice(time_3to4)
        while Ipad2_3to4 == Ipad1_2to3 or Ipad2_3to4 == Ipad2_2to3 or Ipad2_3to4 == Ipad3_2to3:   
            if len(time_3to4) == 1:
                Ipad2_3to4 = ()   
                break
            else:
                Ipad2_3to4 = random.choice(time_3to4)
        time_3to4.remove(Ipad2_3to4)
        if time_3to4 == []:
            continue
    #Desserts
        Desserts3to4 = random.choice(time_3to4)
        time_3to4.remove(Desserts3to4)
        if time_3to4 == []:
            continue
    #Register 2
        Register2_3to4 = random.choice(time_3to4)
        time_3to4.remove(Register2_3to4)
        if time_3to4 == []:
            continue
    #IPad 3
        Ipad3_3to4 = random.choice(time_3to4)
        while Ipad3_3to4 == Ipad1_2to3 or Ipad3_3to4 == Ipad2_2to3 or Ipad3_3to4 == Ipad3_2to3:
                if len(time_3to4) == 1:
                    Ipad3_3to4 = ()   
                    break
                else:
                    Ipad3_3to4 = random.choice(time_3to4)
        time_3to4.remove(Ipad3_3to4)
        if time_3to4 == []:
            continue
    #Dining Room
        Dining3to4 = random.choice(time_3to4)
        time_3to4.remove(Dining3to4)
        if time_3to4 == []:
            continue
    #Sauces
        Sauce3to4 = random.choice(time_3to4)
        time_3to4.remove(Sauce3to4)
        if time_3to4 == []:
            break
#4 to 5
for x in time_4to5:
    #Register 1
        Register1_4to5 = random.choice(time_4to5)
        time_4to5.remove(Register1_4to5)
        if time_4to5 == []:
            continue          
    #Window    
        window4to5 = random.choice(time_4to5)
        time_4to5.remove(window4to5)
        if time_4to5 == []:
            continue
    #DT bagging    
        DTbag4to5 = random.choice(time_4to5)
        time_4to5.remove(DTbag4to5)
        if time_4to5 == []:
            continue
    #Drinks
        Drinks4to5 = random.choice(time_4to5)
        time_4to5.remove(Drinks4to5)
        if time_4to5 == []:
            continue
    #IPad 1
        Ipad1_4to5 = ()
        Ipad2_4to5 = ()
        Ipad3_4to5 = ()
        Ipad1_4to5 = random.choice(time_4to5)
        while Ipad1_4to5 == Ipad1_3to4 or Ipad1_4to5 == Ipad2_3to4 or Ipad1_4to5 == Ipad3_3to4:
                if len(time_4to5) == 1:
                    Ipad1_4to5 = ()   
                    break
                else:
                    Ipad1_4to5 = random.choice(time_4to5)
        time_4to5.remove(Ipad1_4to5)
        if time_4to5 == []:
            continue
    #Cash
        Cash4to5 = random.choice(time_4to5)
        time_4to5.remove(Cash4to5)
        if time_4to5 == []:
            continue
    #FC bagging
        FCbag4to5 = random.choice(time_4to5)
        time_4to5.remove(FCbag4to5)
        if time_4to5 == []:
            continue
    #IPad 2
        Ipad2_4to5 = random.choice(time_4to5)
        while Ipad2_4to5 == Ipad1_3to4 or Ipad2_4to5 == Ipad2_3to4 or Ipad2_4to5 == Ipad3_3to4:   
            if len(time_4to5) == 1:
                Ipad2_4to5 = ()   
                break
            else:
                Ipad2_4to5 = random.choice(time_4to5)
        time_4to5.remove(Ipad2_4to5)
        if time_4to5 == []:
            continue
    #Desserts
        Desserts4to5 = random.choice(time_4to5)
        time_4to5.remove(Desserts4to5)
        if time_4to5 == []:
            continue
    #Register 2
        Register2_4to5 = random.choice(time_4to5)
        time_4to5.remove(Register2_4to5)
        if time_4to5 == []:
            continue
    #IPad 3
        Ipad3_4to5 = random.choice(time_4to5)
        while Ipad3_4to5 == Ipad1_3to4 or Ipad3_4to5 == Ipad2_3to4 or Ipad3_4to5 == Ipad3_3to4:  
            if len(time_4to5) == 1:
                Ipad3_4to5 = ()    
                break
            else:
                Ipad3_4to5 = random.choice(time_4to5)
        time_4to5.remove(Ipad3_4to5)
        if time_4to5 == []:
            continue
    #Dining Room
        Dining4to5 = random.choice(time_4to5)
        time_4to5.remove(Dining4to5)
        if time_4to5 == []:
            continue
    #Sauces
        Sauce4to5 = random.choice(time_4to5)
        time_4to5.remove(Sauce4to5)
        if time_4to5 == []:
            break
#5 to 6
for x in time_5to6:
    #Register 1
        Register1_5to6 = random.choice(time_5to6)
        time_5to6.remove(Register1_5to6)
        if time_5to6 == []:
            continue          
    #Window    
        window5to6 = random.choice(time_5to6)
        time_5to6.remove(window5to6)
        if time_5to6 == []:
            continue
    #DT bagging    
        DTbag5to6 = random.choice(time_5to6)
        time_5to6.remove(DTbag5to6)
        if time_5to6 == []:
            continue
    #Drinks
        Drinks5to6 = random.choice(time_5to6)
        time_5to6.remove(Drinks5to6)
        if time_5to6 == []:
            continue
    #IPad 1
        Ipad1_5to6 = ()
        Ipad2_5to6 = ()
        Ipad3_5to6 = ()
        Ipad1_5to6 = random.choice(time_5to6)
        while Ipad1_5to6 == Ipad1_4to5 or Ipad1_5to6 == Ipad2_4to5 or Ipad1_5to6 == Ipad3_4to5:
            if len(time_5to6) == 1:
                Ipad1_5to6 = ()   
                break
            else:
                Ipad1_5to6 = random.choice(time_5to6)
        time_5to6.remove(Ipad1_5to6)
        if time_5to6 == []:
            continue
    #Cash
        Cash5to6 = random.choice(time_5to6)
        time_5to6.remove(Cash5to6)
        if time_5to6 == []:
            continue
    #FC bagging
        FCbag5to6 = random.choice(time_5to6)
        time_5to6.remove(FCbag5to6)
        if time_5to6 == []:
            continue
    #IPad 2
        Ipad2_5to6 = random.choice(time_5to6)
        while Ipad2_5to6 == Ipad1_4to5 or Ipad2_5to6 == Ipad2_4to5 or Ipad2_5to6 == Ipad3_4to5:   
            if len(time_5to6) == 1:
                Ipad2_5to6 = ()   
                break
            else:
                Ipad2_5to6 = random.choice(time_5to6)
        time_5to6.remove(Ipad2_5to6)
        if time_5to6 == []:
            continue
    #Desserts
        Desserts5to6 = random.choice(time_5to6)
        time_5to6.remove(Desserts5to6)
        if time_5to6 == []:
            continue
    #Register 2
        Register2_5to6 = random.choice(time_5to6)
        time_5to6.remove(Register2_5to6)
        if time_5to6 == []:
            continue
    #IPad 3
        Ipad3_5to6 = random.choice(time_5to6)
        while Ipad3_5to6 == Ipad1_4to5 or Ipad3_5to6 == Ipad2_4to5 or Ipad3_5to6 == Ipad3_4to5:
                if len(time_5to6) == 1:
                    Ipad3_5to6 = ()   
                    break
                else:
                    Ipad3_5to6 = random.choice(time_5to6)
        time_5to6.remove(Ipad3_5to6)
        if time_5to6 == []:
            continue
    #Dining Room
        Dining5to6 = random.choice(time_5to6)
        time_5to6.remove(Dining5to6)
        if time_5to6 == []:
            continue
    #Sauces
        Sauce5to6 = random.choice(time_5to6)
        time_5to6.remove(Sauce5to6)
        if time_5to6 == []:
            break
#6 to 7
for x in time_6to7:
    #Register 1
        Register1_6to7 = random.choice(time_6to7)
        time_6to7.remove(Register1_6to7)
        if time_6to7 == []:
            continue          
    #Window    
        window6to7 = random.choice(time_6to7)
        time_6to7.remove(window6to7)
        if time_6to7 == []:
            continue
    #DT bagging    
        DTbag6to7 = random.choice(time_6to7)
        time_6to7.remove(DTbag6to7)
        if time_6to7 == []:
            continue
    #Drinks
        Drinks6to7 = random.choice(time_6to7)
        time_6to7.remove(Drinks6to7)
        if time_6to7 == []:
            continue
    #IPad 1
        Ipad1_6to7 = ()
        Ipad2_6to7 = ()
        Ipad3_6to7 = ()
        Ipad1_6to7 = random.choice(time_6to7)
        while Ipad1_6to7 == Ipad1_5to6 or Ipad1_6to7 == Ipad2_5to6 or Ipad1_6to7 == Ipad3_5to6:
            if len(time_6to7) == 1:
                Ipad1_6to7 = ()   
                break
            else:
                Ipad1_6to7 = random.choice(time_6to7)           
        time_6to7.remove(Ipad1_6to7)
        if time_6to7 == []:
            continue
    #Cash
        Cash6to7 = random.choice(time_6to7)
        time_6to7.remove(Cash6to7)
        if time_6to7 == []:
            continue
    #FC bagging
        FCbag6to7 = random.choice(time_6to7)
        time_6to7.remove(FCbag6to7)
        if time_6to7 == []:
            continue
    #IPad 2
        Ipad2_6to7 = random.choice(time_6to7)
        while Ipad2_6to7 == Ipad1_5to6 or Ipad2_6to7 == Ipad2_5to6 or Ipad2_6to7 == Ipad3_5to6:
            if len(time_6to7) == 1:
                Ipad2_6to7 = ()   
                break
            else:
                Ipad2_6to7 = random.choice(time_6to7)   
        time_6to7.remove(Ipad2_6to7)
        if time_6to7 == []:
            continue
    #Desserts
        Desserts6to7 = random.choice(time_6to7)
        time_6to7.remove(Desserts6to7)
        if time_6to7 == []:
            continue
    #Register 2
        Register2_6to7 = random.choice(time_6to7)
        time_6to7.remove(Register2_6to7)
        if time_6to7 == []:
            continue
    #IPad 3
        Ipad3_6to7 = random.choice(time_6to7)
        while Ipad3_6to7 == Ipad1_5to6 or Ipad3_6to7 == Ipad2_5to6 or Ipad3_6to7 == Ipad3_5to6:
            if len(time_6to7) == 1:
                Ipad3_6to7 = ()   
                break   
            else:
                Ipad3_6to7 = random.choice(time_6to7)     
        time_6to7.remove(Ipad3_6to7)
        if time_6to7 == []:
            continue
    #Dining Room
        Dining6to7 = random.choice(time_6to7)
        time_6to7.remove(Dining6to7)
        if time_6to7 == []:
            continue
    #Sauces
        Sauce6to7 = random.choice(time_6to7)
        time_6to7.remove(Sauce6to7)
        if time_6to7 == []:
            break
        
#7 to 8
for x in time_7to8:
    #Register 1
        Register1_7to8 = random.choice(time_7to8)
        time_7to8.remove(Register1_7to8)
        if time_7to8 == []:
            continue          
    #Window    
        window7to8 = random.choice(time_7to8)
        time_7to8.remove(window7to8)
        if time_7to8 == []:
            continue
    #DT bagging    
        DTbag7to8 = random.choice(time_7to8)
        time_7to8.remove(DTbag7to8)
        if time_7to8 == []:
            continue
    #Drinks
        Drinks7to8 = random.choice(time_7to8)
        time_7to8.remove(Drinks7to8)
        if time_7to8 == []:
            continue
    #IPad 1
        Ipad1_7to8 = ()
        Ipad2_7to8 = ()
        Ipad3_7to8 = ()
        Ipad1_7to8 = random.choice(time_7to8)
        while Ipad1_7to8 == Ipad1_6to7 or Ipad1_7to8 == Ipad2_6to7 or Ipad1_7to8 == Ipad3_6to7:
            if len(time_7to8) == 1:
                Ipad1_7to8 = ()   
                break   
            else:
                Ipad1_7to8 = random.choice(time_7to8)
        time_7to8.remove(Ipad1_7to8)
        if time_7to8 == []:
            continue
    #Cash
        Cash7to8 = random.choice(time_7to8)
        time_7to8.remove(Cash7to8)
        if time_7to8 == []:
            continue
    #Register 1
        Register1_7to8 = random.choice(time_7to8)
        time_7to8.remove(Register1_7to8)
        if time_7to8 == []:
             continue
    #FC bagging
        FCbag7to8 = random.choice(time_7to8)
        time_7to8.remove(FCbag7to8)
        if time_7to8 == []:
            continue
    #IPad 2
        Ipad2_7to8 = random.choice(time_7to8)
        while Ipad2_7to8 == Ipad1_6to7 or Ipad2_7to8 == Ipad2_6to7 or Ipad2_7to8 == Ipad3_6to7:
            if len(time_7to8) == 1:
                Ipad2_7to8 = ()   
                break   
            else:
                Ipad2_7to8 = random.choice(time_7to8)
        time_7to8.remove(Ipad2_7to8)
        if time_7to8 == []:
            continue
    #Desserts
        Desserts7to8 = random.choice(time_7to8)
        time_7to8.remove(Desserts7to8)
        if time_7to8 == []:
            continue
    #Register 2
        Register2_7to8 = random.choice(time_7to8)
        time_7to8.remove(Register2_7to8)
        if time_7to8 == []:
            continue
    #IPad 3
        Ipad3_7to8 = random.choice(time_7to8)
        while Ipad3_7to8 == Ipad1_6to7 or Ipad3_7to8 == Ipad2_6to7 or Ipad3_7to8 == Ipad3_6to7:
            if len(time_7to8) == 1:
                Ipad3_7to8 = ()   
                break   
            else:
                Ipad3_7to8 = random.choice(time_7to8)
        time_7to8.remove(Ipad3_7to8)
        if time_7to8 == []:
            continue
    #Dining Room
        Dining7to8 = random.choice(time_7to8)
        time_7to8.remove(Dining7to8)
        if time_7to8 == []:
            continue
    #Sauces
        Sauce7to8 = random.choice(time_7to8)
        time_7to8.remove(Sauce7to8)
        if time_7to8 == []:
            break
#8 to 9            
for x in time_8to9:
#Window    
        window8to9 = random.choice(time_8to9)
        time_8to9.remove(window8to9)
        if time_8to9 == []:
            continue
    #DT bagging    
        DTbag8to9 = random.choice(time_8to9)
        time_8to9.remove(DTbag8to9)
        if time_8to9 == []:
            continue
    #Drinks
        Drinks8to9 = random.choice(time_8to9)
        time_8to9.remove(Drinks8to9)
        if time_8to9 == []:
            continue
    #Dining Room
        Dining8to9 = random.choice(time_8to9)
        time_8to9.remove(Dining8to9)
        if time_8to9 == []:
            continue        
    #IPad 1
        Ipad1_8to9 = ()
        Ipad2_8to9 = ()
        Ipad3_8to9 = ()
        Ipad1_8to9 = random.choice(time_8to9)
        while Ipad1_8to9 == Ipad1_7to8 or Ipad1_8to9 == Ipad2_7to8 or Ipad1_8to9 == Ipad3_7to8:
            if len(time_8to9) == 1:
                Ipad1_8to9 = ()   
                break  
            else:
                Ipad1_8to9 = random.choice(time_8to9)
        time_8to9.remove(Ipad1_8to9)
        if time_8to9 == []:
            continue
    #Cash
        Cash8to9 = random.choice(time_8to9)
        time_8to9.remove(Cash8to9)
        if time_8to9 == []:
            continue
    #Register 1
        Register1_8to9 = random.choice(time_8to9)
        time_8to9.remove(Register1_8to9)
        if time_8to9 == []:
             continue
    #FC bagging
        FCbag8to9 = random.choice(time_8to9)
        time_8to9.remove(FCbag8to9)
        if time_8to9 == []:
            continue
    #IPad 2
        Ipad2_8to9 = random.choice(time_8to9)
        while Ipad2_8to9 == Ipad1_7to8 or Ipad2_8to9 == Ipad2_7to8 or Ipad2_8to9 == Ipad3_7to8:
            if len(time_8to9) == 1:
                Ipad2_8to9 = ()   
                break  
            else:
                Ipad2_8to9 = random.choice(time_8to9)
        time_8to9.remove(Ipad2_8to9)
        if time_8to9 == []:
            continue
    #Desserts
        Desserts8to9 = random.choice(time_8to9)
        time_8to9.remove(Desserts8to9)
        if time_8to9 == []:
            continue
    #Register 2
        Register2_8to9 = random.choice(time_8to9)
        time_8to9.remove(Register2_8to9)
        if time_8to9 == []:
            continue
    #IPad 3
        Ipad3_8to9 = random.choice(time_8to9)
        while Ipad3_8to9 == Ipad1_7to8 or Ipad3_8to9 == Ipad2_7to8 or Ipad3_8to9 == Ipad3_7to8:
            if len(time_8to9) == 1:
                Ipad3_8to9 = ()   
                break  
            else:
                Ipad3_8to9 = random.choice(time_8to9)
        time_8to9.remove(Ipad3_8to9)
        if time_8to9 == []:
            continue
    #Sauces
        Sauce8to9 = random.choice(time_8to9)
        time_8to9.remove(Sauce8to9)
        if time_8to9 == []:
            break        
#9 to 10            
for x in time_9to10:
     #Window    
        window9to10 = random.choice(time_9to10)
        time_9to10.remove(window9to10)
        if time_9to10 == []:
            continue
    #DT bagging    
        DTbag9to10 = random.choice(time_9to10)
        time_9to10.remove(DTbag9to10)
        if time_9to10 == []:
            continue
    #Drinks
        Drinks9to10 = random.choice(time_9to10)
        time_9to10.remove(Drinks9to10)
        if time_9to10 == []:
            continue
    #Dining Room
        Dining9to10 = random.choice(time_9to10)
        time_9to10.remove(Dining9to10)
        if time_9to10 == []:
            continue
    #Register 1
        Register1_9to10 = random.choice(time_9to10)
        time_9to10.remove(Register1_9to10)
        if time_9to10 == []:
             continue
    #FC bagging
        FCbag9to10 = random.choice(time_9to10)
        time_9to10.remove(FCbag9to10)
        if time_9to10 == []:
            continue

#Output as word document

def iter_paragraphs_of_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                  yield paragraph
                yield from iter_paragraphs_of_tables(cell.tables)

for paragraph in iter_paragraphs_of_tables(doc.tables):
     if '{{Win2to3}}' in paragraph.text:
      if window2to3 != ():
       paragraph.text = paragraph.text.replace("{{Win2to3}}", window2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Win2to3}}", "")   
     if '{{DTBag2to3}}' in paragraph.text:
      if DTbag2to3 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", DTbag2to3 + " 2-3")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag2to3}}", "")
     if '{{Drinks2to3}}' in paragraph.text:
      if Drinks2to3 != ():
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", Drinks2to3 + " 2-3")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks2to3}}", "")   
     if '{{Desserts2to3}}' in paragraph.text:
        if Desserts2to3 != ():
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", Desserts2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts2to3}}", "")   
     if '{{Sauce2to3}}' in paragraph.text:
        if Sauce2to3 != ():
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", Sauce2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce2to3}}", "")   
     if '{{FCBag2to3}}' in paragraph.text:
        if FCbag2to3 != ():
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", FCbag2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag2to3}}", "")   
     if '{{Reg1_2to3}}' in paragraph.text:
        if Register1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", Register1_2to3 + " 2-3")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_2to3}}", "") 
     if '{{Reg2_2to3}}' in paragraph.text:
        if Register2_2to3 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", Register2_2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_2to3}}", "")  
     if '{{Dining_2to3}}' in paragraph.text:
        if Dining2to3 != ():
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", Dining2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_2to3}}", "")
     if '{{Cash2to3}}' in paragraph.text:
        if Cash2to3 != ():
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", Cash2to3 + " 2-3")
        else:
         paragraph.text = paragraph.text.replace("{{Cash2to3}}", "")   
     if '{{IPad1_2to3}}' in paragraph.text:
        if Ipad1_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", Ipad1_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_2to3}}", "")    
     if '{{IPad2_2to3}}' in paragraph.text:
        if Ipad2_2to3 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", Ipad2_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_2to3}}", "")   
     if '{{IPad3_2to3}}' in paragraph.text:
        if Ipad3_2to3 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", Ipad3_2to3)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_2to3}}", "")
     if '{{Win3to4}}' in paragraph.text:
      if window3to4 != ():
       paragraph.text = paragraph.text.replace("{{Win3to4}}", window3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Win3to4}}", "")   
     if '{{DTBag3to4}}' in paragraph.text:
      if DTbag3to4 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", DTbag3to4 + " 3-4")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag3to4}}", "")
     if '{{Drinks3to4}}' in paragraph.text:
      if Drinks3to4 != ():
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", Drinks3to4 + " 3-4")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks3to4}}", "")   
     if '{{Desserts3to4}}' in paragraph.text:
        if Desserts3to4 != ():
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", Desserts3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts3to4}}", "")   
     if '{{Sauce3to4}}' in paragraph.text:
        if Sauce3to4 != ():
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", Sauce3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce3to4}}", "")   
     if '{{FCBag3to4}}' in paragraph.text:
        if FCbag3to4 != ():
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", FCbag3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag3to4}}", "")   
     if '{{Reg1_3to4}}' in paragraph.text:
        if Register1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", Register1_3to4 + " 3-4")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_3to4}}", "") 
     if '{{Reg2_3to4}}' in paragraph.text:
        if Register2_3to4 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", Register2_3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_3to4}}", "")  
     if '{{Dining_3to4}}' in paragraph.text:
        if Dining3to4 != ():
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", Dining3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_3to4}}", "")
     if '{{Cash3to4}}' in paragraph.text:
        if Cash3to4 != ():
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", Cash3to4 + " 3-4")
        else:
         paragraph.text = paragraph.text.replace("{{Cash3to4}}", "")   
     if '{{IPad1_3to4}}' in paragraph.text:
        if Ipad1_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", Ipad1_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_3to4}}", "")    
     if '{{IPad2_3to4}}' in paragraph.text:
        if Ipad2_3to4 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", Ipad2_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_3to4}}", "")   
     if '{{IPad3_3to4}}' in paragraph.text:
        if Ipad3_3to4 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", Ipad3_3to4)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_3to4}}", "")       
     if '{{Win4to5}}' in paragraph.text:
      if window4to5 != ():
       paragraph.text = paragraph.text.replace("{{Win4to5}}", window4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Win4to5}}", "")   
     if '{{DTBag4to5}}' in paragraph.text:
      if DTbag4to5 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", DTbag4to5 + " 4-5")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag4to5}}", "")
     if '{{Drinks4to5}}' in paragraph.text:
      if Drinks4to5 != ():
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", Drinks4to5 + " 4-5")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks4to5}}", "")   
     if '{{Desserts4to5}}' in paragraph.text:
        if Desserts4to5 != ():
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", Desserts4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts4to5}}", "")   
     if '{{Sauce4to5}}' in paragraph.text:
        if Sauce4to5 != ():
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", Sauce4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce4to5}}", "")   
     if '{{FCBag4to5}}' in paragraph.text:
        if FCbag4to5 != ():
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", FCbag4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag4to5}}", "")   
     if '{{Reg1_4to5}}' in paragraph.text:
        if Register1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", Register1_4to5 + " 4-5")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_4to5}}", "") 
     if '{{Reg2_4to5}}' in paragraph.text:
        if Register2_4to5 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", Register2_4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_4to5}}", "")  
     if '{{Dining_4to5}}' in paragraph.text:
        if Dining4to5 != ():
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", Dining4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_4to5}}", "")
     if '{{Cash4to5}}' in paragraph.text:
        if Cash4to5 != ():
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", Cash4to5 + " 4-5")
        else:
         paragraph.text = paragraph.text.replace("{{Cash4to5}}", "")   
     if '{{IPad1_4to5}}' in paragraph.text:
        if Ipad1_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", Ipad1_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_4to5}}", "")    
     if '{{IPad2_4to5}}' in paragraph.text:
        if Ipad2_4to5 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", Ipad2_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_4to5}}", "")   
     if '{{IPad3_4to5}}' in paragraph.text:
        if Ipad3_4to5 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", Ipad3_4to5)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_4to5}}", "")
     if '{{Win5to6}}' in paragraph.text:
      if window5to6 != ():
       paragraph.text = paragraph.text.replace("{{Win5to6}}", window5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Win5to6}}", "")   
     if '{{DTBag5to6}}' in paragraph.text:
      if DTbag5to6 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", DTbag5to6 + " 5-6")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag5to6}}", "")
     if '{{Drinks5to6}}' in paragraph.text:
      if Drinks5to6 != ():
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", Drinks5to6 + " 5-6")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks5to6}}", "")   
     if '{{Desserts5to6}}' in paragraph.text:
        if Desserts5to6 != ():
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", Desserts5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts5to6}}", "")   
     if '{{Sauce5to6}}' in paragraph.text:
        if Sauce5to6 != ():
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", Sauce5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce5to6}}", "")   
     if '{{FCBag5to6}}' in paragraph.text:
        if FCbag5to6 != ():
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", FCbag5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag5to6}}", "")   
     if '{{Reg1_5to6}}' in paragraph.text:
        if Register1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", Register1_5to6 + " 5-6")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_5to6}}", "") 
     if '{{Reg2_5to6}}' in paragraph.text:
        if Register2_5to6 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", Register2_5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_5to6}}", "")  
     if '{{Dining_5to6}}' in paragraph.text:
        if Dining5to6 != ():
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", Dining5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_5to6}}", "")
     if '{{Cash5to6}}' in paragraph.text:
        if Cash5to6 != ():
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", Cash5to6 + " 5-6")
        else:
         paragraph.text = paragraph.text.replace("{{Cash5to6}}", "")   
     if '{{IPad1_5to6}}' in paragraph.text:
        if Ipad1_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", Ipad1_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_5to6}}", "")    
     if '{{IPad2_5to6}}' in paragraph.text:
        if Ipad2_5to6 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", Ipad2_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_5to6}}", "")   
     if '{{IPad3_5to6}}' in paragraph.text:
        if Ipad3_5to6 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", Ipad3_5to6)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_5to6}}", "")         
     if '{{Win6to7}}' in paragraph.text:
      if window6to7 != ():
       paragraph.text = paragraph.text.replace("{{Win6to7}}", window6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Win6to7}}", "")   
     if '{{DTBag6to7}}' in paragraph.text:
      if DTbag6to7 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", DTbag6to7 + " 6-7")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag6to7}}", "")
     if '{{Drinks6to7}}' in paragraph.text:
      if Drinks6to7 != ():
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", Drinks6to7 + " 6-7")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks6to7}}", "")   
     if '{{Desserts6to7}}' in paragraph.text:
        if Desserts6to7 != ():
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", Desserts6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts6to7}}", "")   
     if '{{Sauce6to7}}' in paragraph.text:
        if Sauce6to7 != ():
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", Sauce6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce6to7}}", "")   
     if '{{FCBag6to7}}' in paragraph.text:
        if FCbag6to7 != ():
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", FCbag6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag6to7}}", "")   
     if '{{Reg1_6to7}}' in paragraph.text:
        if Register1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", Register1_6to7 + " 6-7")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_6to7}}", "") 
     if '{{Reg2_6to7}}' in paragraph.text:
        if Register2_6to7 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", Register2_6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_6to7}}", "")  
     if '{{Dining_6to7}}' in paragraph.text:
        if Dining6to7 != ():
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", Dining6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_6to7}}", "")
     if '{{Cash6to7}}' in paragraph.text:
        if Cash6to7 != ():
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", Cash6to7 + " 6-7")
        else:
         paragraph.text = paragraph.text.replace("{{Cash6to7}}", "")   
     if '{{IPad1_6to7}}' in paragraph.text:
        if Ipad1_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", Ipad1_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_6to7}}", "")    
     if '{{IPad2_6to7}}' in paragraph.text:
        if Ipad2_6to7 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", Ipad2_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_6to7}}", "")   
     if '{{IPad3_6to7}}' in paragraph.text:
        if Ipad3_6to7 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", Ipad3_6to7)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_6to7}}", "")    
     if '{{Win7to8}}' in paragraph.text:
      if window7to8 != ():
       paragraph.text = paragraph.text.replace("{{Win7to8}}", window7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Win7to8}}", "")   
     if '{{DTBag7to8}}' in paragraph.text:
      if DTbag7to8 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", DTbag7to8 + " 7-8")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag7to8}}", "")
     if '{{Drinks7to8}}' in paragraph.text:
      if Drinks7to8 != ():
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", Drinks7to8 + " 7-8")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks7to8}}", "")   
     if '{{Desserts7to8}}' in paragraph.text:
        if Desserts7to8 != ():
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", Desserts7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts7to8}}", "")   
     if '{{Sauce7to8}}' in paragraph.text:
        if Sauce7to8 != ():
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", Sauce7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce7to8}}", "")   
     if '{{FCBag7to8}}' in paragraph.text:
        if FCbag7to8 != ():
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", FCbag7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag7to8}}", "")   
     if '{{Reg1_7to8}}' in paragraph.text:
        if Register1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", Register1_7to8 + " 7-8")
     else:
        paragraph.text = paragraph.text.replace("{{Reg1_7to8}}", "") 
     if '{{Reg2_7to8}}' in paragraph.text:
        if Register2_7to8 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", Register2_7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_7to8}}", "")  
     if '{{Dining_7to8}}' in paragraph.text:
        if Dining7to8 != ():
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", Dining7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_7to8}}", "")
     if '{{Cash7to8}}' in paragraph.text:
        if Cash7to8 != ():
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", Cash7to8 + " 7-8")
        else:
         paragraph.text = paragraph.text.replace("{{Cash7to8}}", "")   
     if '{{IPad1_7to8}}' in paragraph.text:
        if Ipad1_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", Ipad1_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_7to8}}", "")    
     if '{{IPad2_7to8}}' in paragraph.text:
        if Ipad2_7to8 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", Ipad2_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_7to8}}", "")   
     if '{{IPad3_7to8}}' in paragraph.text:
        if Ipad3_7to8 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", Ipad3_7to8)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_7to8}}", "")
     if '{{Win8to9}}' in paragraph.text:
      if window8to9 != ():
       paragraph.text = paragraph.text.replace("{{Win8to9}}", window8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Win8to9}}", "")   
     if '{{DTBag8to9}}' in paragraph.text:
      if DTbag8to9 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", DTbag8to9 + " 8-9")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag8to9}}", "")
     if '{{Drinks8to9}}' in paragraph.text:
      if Drinks8to9 != ():
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", Drinks8to9 + " 8-9")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks8to9}}", "")   
     if '{{Desserts8to9}}' in paragraph.text:
        if Desserts8to9 != ():
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", Desserts8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts8to9}}", "")   
     if '{{Sauce8to9}}' in paragraph.text:
        if Sauce8to9 != ():
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", Sauce8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce8to9}}", "")   
     if '{{FCBag8to9}}' in paragraph.text:
        if FCbag8to9 != ():
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", FCbag8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag8to9}}", "")   
     if '{{Reg1_8to9}}' in paragraph.text:
        if Register1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", Register1_8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Reg1_8to9}}", "") 
     if '{{Reg2_8to9}}' in paragraph.text:
        if Register2_8to9 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", Register2_8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_8to9}}", "")  
     if '{{Dining_8to9}}' in paragraph.text:
        if Dining8to9 != ():
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", Dining8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_8to9}}", "")
     if '{{Cash8to9}}' in paragraph.text:
        if Cash8to9 != ():
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", Cash8to9 + " 8-9")
        else:
         paragraph.text = paragraph.text.replace("{{Cash8to9}}", "")   
     if '{{IPad1_8to9}}' in paragraph.text:
        if Ipad1_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", Ipad1_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad1_8to9}}", "")    
     if '{{IPad2_8to9}}' in paragraph.text:
        if Ipad2_8to9 != (): 
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", Ipad2_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad2_8to9}}", "")   
     if '{{IPad3_8to9}}' in paragraph.text:
        if Ipad3_8to9 != ():
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", Ipad3_8to9)
        else:
         paragraph.text = paragraph.text.replace("{{IPad3_8to9}}", "")         
     if '{{Win9to10}}' in paragraph.text:
      if window9to10 != ():
       paragraph.text = paragraph.text.replace("{{Win9to10}}", window9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Win9to10}}", "")   
     if '{{DTBag9to10}}' in paragraph.text:
      if DTbag9to10 != ():  
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", DTbag9to10 + " 9-10")
      else:
         paragraph.text = paragraph.text.replace("{{DTBag9to10}}", "")
     if '{{Drinks9to10}}' in paragraph.text:
      if Drinks9to10 != ():
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", Drinks9to10 + " 9-10")
      else:
       paragraph.text = paragraph.text.replace("{{Drinks9to10}}", "")   
     if '{{Desserts9to10}}' in paragraph.text:
        if Desserts9to10 != ():
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", Desserts9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Desserts9to10}}", "")   
     if '{{Sauce9to10}}' in paragraph.text:
        if Sauce9to10 != ():
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", Sauce9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Sauce9to10}}", "")   
     if '{{FCBag9to10}}' in paragraph.text:
        if FCbag9to10 != ():
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", FCbag9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{FCBag9to10}}", "")   
     if '{{Reg1_9to10}}' in paragraph.text:
        if Register1_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", Register1_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg1_9to10}}", "") 
     if '{{Reg2_9to10}}' in paragraph.text:
        if Register2_9to10 != ():
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", Register2_9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Reg2_9to10}}", "")  
     if '{{Dining_9to10}}' in paragraph.text:
        if Dining9to10 != ():
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", Dining9to10 + " 9-10")
        else:
         paragraph.text = paragraph.text.replace("{{Dining_9to10}}", "")
     for y in range (0,len (FOHCrew)):
         team_order = '{{FOHCrew'+ str (y) + '}}'
         if team_order in paragraph.text:
             paragraph.text = paragraph.text.replace(team_order, FOHCrew[y])
if '{{Date}}' in paragraph.text:
  paragraph.text = paragraph.text.replace("{{Date}}", print_date)            
font = paragraph.style.font         
font.size = Pt(6.5)         
doc.save('ShiftPlanner.docx')         
os.startfile('C:/Users/Brand/Documents/ShiftPlanner.docx')

